"""
Advanced Document Processor for Surveyors, Engineers, and Geomaticians

This module provides atomic, AI-driven document extraction tools that allow
the agent to dynamically extract information based on the document type and task.

Key Philosophy:
- Provide low-level atomic tools (get_text, get_tables, get_metadata)
- Let the AI agent reason about what to extract
- No hardcoded extraction patterns - agent decides what's relevant
- Support complex document structures (tables, signatures, headers/footers)
"""

import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
import docx
import PyPDF2
import pdfplumber
import zipfile
import xml.etree.ElementTree as ET
from collections import deque
import html
from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """
    Advanced document processor with atomic extraction methods.
    
    This class provides low-level tools that the AI agent can use dynamically
    to extract information from documents. Instead of hardcoded patterns,
    the agent reasons about what to extract based on the document type and task.
    
    Usage Pattern:
        The agent should:
        1. Get document metadata to understand structure
        2. Get full text for analysis
        3. Get tables if present
        4. Extract specific information using reasoning
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self._word_cache: Dict[str, Any] = {}
        self._pdf_cache: Dict[str, Any] = {}

    # ==========================================================================
    # FAST WORD STATS (docProps/app.xml)
    # ==========================================================================

    def _get_word_docprops_stats(self, file_path: str) -> Dict[str, Optional[int]]:
        """
        Read Word's internal document statistics (fast).
        
        This avoids parsing the full document body and is typically accurate even
        when most text lives inside tables.
        
        Returns:
            Dict with: pages, words, characters (values may be None)
        """
        try:
            with zipfile.ZipFile(file_path) as zf:
                app_xml = zf.read("docProps/app.xml")
            root = ET.fromstring(app_xml)
            ns = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}

            def _to_int(node):
                if node is None or node.text is None:
                    return None
                s = node.text.strip()
                return int(s) if s.isdigit() else None

            pages = _to_int(root.find("ep:Pages", ns))
            words = _to_int(root.find("ep:Words", ns))
            chars = _to_int(root.find("ep:Characters", ns))
            return {"pages": pages, "words": words, "characters": chars}
        except Exception:
            return {"pages": None, "words": None, "characters": None}

    def _iter_docx_lines_from_xml(self, file_path: str):
        """
        Yield approximate text 'lines' from a .docx by scanning word/document.xml.
        
        This is much faster than python-docx for very large / table-heavy documents.
        It is a best-effort extraction (formatting may be lost).
        """
        try:
            with zipfile.ZipFile(file_path) as zf:
                xml_bytes = zf.read("word/document.xml")
        except Exception as e:
            logger.warning(f"Could not read docx XML: {e}")
            return

        try:
            xml_text = xml_bytes.decode("utf-8", errors="ignore")
        except Exception:
            xml_text = str(xml_bytes)

        # Add paragraph boundaries so we can emit grouped lines.
        xml_text = xml_text.replace("</w:p>", "\n")

        # Replace common Word break tags with newlines
        xml_text = re.sub(r"<w:br\s*/>", "\n", xml_text, flags=re.IGNORECASE)
        xml_text = re.sub(r"<w:tab\s*/>", "\t", xml_text, flags=re.IGNORECASE)

        # Extract text per paragraph by splitting on newlines in the XML
        for para_xml in xml_text.split("\n"):
            parts = []
            for raw in re.findall(r"<w:t[^>]*>(.*?)</w:t>", para_xml, flags=re.IGNORECASE | re.DOTALL):
                parts.append(html.unescape(re.sub(r"<[^>]+>", "", raw)))
            line = "".join(parts).strip()
            if line:
                yield line
    
    # ==========================================================================
    # ATOMIC EXTRACTION METHODS (AI-driven)
    # ==========================================================================
    
    def get_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Get basic metadata about the document.
        
        This is an atomic method that returns document structure information
        without extracting content. The agent can use this to understand
        what extraction methods to use.
        
        Args:
            file_path: Path to document
            
        Returns:
            Dict containing:
            - file_path: Full path
            - file_type: "Word" or "PDF"
            - file_size: Size in bytes
            - page_count: Number of pages (for PDF) or estimated pages (for Word)
            - has_tables: Whether document contains tables
            - creation_date: File creation date
            - modification_date: File modification date
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}
        
        metadata = {
            "success": True,
            "file_path": str(path.resolve()),
            "file_name": path.name,
            "file_type": "Word" if path.suffix.lower() in ['.docx', '.doc'] else "PDF" if path.suffix.lower() == '.pdf' else "Unknown",
            "file_size": path.stat().st_size,
            "creation_date": datetime.fromtimestamp(path.stat().st_ctime).isoformat(),
            "modification_date": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        }
        
        try:
            if metadata["file_type"] == "Word":
                # Prefer Word's internal docProps stats (accurate even when text is in tables)
                stats = self._get_word_docprops_stats(str(path))

                doc = docx.Document(path)

                metadata["has_tables"] = len(doc.tables) > 0
                metadata["table_count"] = len(doc.tables)
                metadata["paragraph_count"] = len(doc.paragraphs)

                # Fallback estimates from paragraphs only (used only if docProps stats unavailable)
                para_word_count = sum(len(para.text.split()) for para in doc.paragraphs)
                para_char_count = sum(len(para.text) for para in doc.paragraphs)

                metadata["page_count"] = stats["pages"] if stats.get("pages") is not None else max(1, para_word_count // 250)
                metadata["word_count"] = stats["words"] if stats.get("words") is not None else para_word_count
                metadata["character_count"] = stats["characters"] if stats.get("characters") is not None else para_char_count

                # Rough token estimate
                if metadata.get("character_count"):
                    metadata["estimated_tokens"] = int(metadata["character_count"]) // 4
                elif metadata.get("word_count"):
                    metadata["estimated_tokens"] = int(int(metadata["word_count"]) / 0.75)
                else:
                    metadata["estimated_tokens"] = 0
            elif metadata["file_type"] == "PDF":
                with pdfplumber.open(path) as pdf:
                    metadata["page_count"] = len(pdf.pages)
                    # Check for tables
                    has_tables = False
                    for page in pdf.pages[:3]:  # Check first 3 pages
                        if page.extract_tables():
                            has_tables = True
                            break
                    metadata["has_tables"] = has_tables
        except Exception as e:
            logger.warning(f"Could not extract full metadata: {e}")
            metadata["metadata_error"] = str(e)
        
        return metadata
    
    def get_full_text(self, file_path: str, preserve_structure: bool = True) -> Dict[str, Any]:
        """
        Extract all text content from the document.
        
        This is an atomic method that returns raw text. The agent can then
        reason about what information to extract from it.
        
        WARNING: For large documents (>100 pages or >50K words), this method
        will automatically check size first and may return a warning instead of
        processing the entire document. Use get_resource_estimation() first
        for large documents, then use extract_sections_by_keywords() instead.
        
        Args:
            file_path: Path to document
            preserve_structure: If True, preserve paragraph breaks and structure
            
        Returns:
            Dict containing:
            - success: Boolean
            - text: Full text content (or warning message for large docs)
            - text_length: Character count
            - page_count: Number of pages
            - structure_preserved: Whether structure was preserved
            - warning: Warning message if document is too large
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}
        
        # Check document size first for large documents (fast check)
        try:
            # Quick file size check first (much faster than reading document)
            file_size_mb = path.stat().st_size / (1024 * 1024)

            # Word: prefer internal stats; file size is not a reliable indicator (docx is compressed)
            if path.suffix.lower() in ['.docx', '.doc']:
                stats = self._get_word_docprops_stats(str(path))
                pages = stats.get("pages")
                words = stats.get("words")
                chars = stats.get("characters")
                estimated_tokens = (chars // 4) if chars else (int(words / 0.75) if words else 0)

                is_large = (
                    (pages is not None and pages > 100) or
                    (words is not None and words > 50000) or
                    (estimated_tokens and estimated_tokens > 100000)
                )
                if is_large:
                    warning_msg = (
                        f"⚠️ LARGE DOCUMENT DETECTED:\n"
                        f"  • File size: {file_size_mb:.2f} MB\n"
                        f"  • Pages (Word stats): {pages if pages is not None else 'unknown'}\n"
                        f"  • Words (Word stats): {words if words is not None else 'unknown'}\n"
                        f"  • Characters (Word stats): {chars if chars is not None else 'unknown'}\n"
                        f"  • Estimated tokens: {estimated_tokens:,}\n\n"
                        f"Full extraction may take a long time and cost significantly more.\n\n"
                        f"RECOMMENDED (cheaper + faster):\n"
                        f"1) document_get_resource_estimation(file_path)\n"
                        f"2) document_get_structure(file_path)\n"
                        f"3) document_extract_sections_by_keywords(file_path, keywords=[...])\n\n"
                        f"If you still want FULL extraction, reply 'yes' and use document_get_text_force(file_path). (yes/no)"
                    )
                    logger.warning(warning_msg)
                    return {
                        "success": False,
                        "error": "Document too large for safe full text extraction",
                        "warning": warning_msg,
                        "page_count": pages,
                        "word_count": words,
                        "character_count": chars,
                        "file_size_mb": round(file_size_mb, 2),
                        "estimated_tokens": estimated_tokens,
                        "recommendation": "Use document_extract_sections_by_keywords() or document_get_text_force() if you confirm"
                    }

            # Non-Word: keep a coarse file-size heuristic
            if path.suffix.lower() == ".pdf" and file_size_mb > 20:
                logger.warning(f"⚠️ Large PDF detected ({file_size_mb:.2f} MB). Consider section extraction.")
        except Exception as e:
            logger.warning(f"Could not check document size: {e}")
            # Continue with extraction if size check fails
        
        try:
            if path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_word_text(path, preserve_structure)
            elif path.suffix.lower() == '.pdf':
                return self._extract_pdf_text(path, preserve_structure)
            else:
                return {"success": False, "error": f"Unsupported format: {path.suffix}"}
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return {"success": False, "error": str(e)}

    def get_full_text_force(self, file_path: str, preserve_structure: bool = True) -> Dict[str, Any]:
        """
        Force full text extraction even for large documents.
        
        Use ONLY after the user explicitly confirms they want full extraction.
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}

        try:
            if path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_word_text(path, preserve_structure)
            elif path.suffix.lower() == '.pdf':
                return self._extract_pdf_text(path, preserve_structure)
            else:
                return {"success": False, "error": f"Unsupported format: {path.suffix}"}
        except Exception as e:
            logger.error(f"Error extracting text (force): {e}")
            return {"success": False, "error": str(e)}
    
    def get_tables(self, file_path: str, page_number: Optional[int] = None) -> Dict[str, Any]:
        """
        Extract all tables from the document.
        
        This is an atomic method that returns tables as structured data.
        The agent can then reason about which tables contain relevant information.
        
        Args:
            file_path: Path to document
            page_number: Optional page number (for PDF) or None for all pages
            
        Returns:
            Dict containing:
            - success: Boolean
            - table_count: Number of tables found
            - tables: List of table dictionaries, each with:
              - page_number: Page where table appears (for PDF)
              - row_count: Number of rows
              - column_count: Number of columns
              - data: 2D array of table data
              - headers: First row if it appears to be headers
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}
        
        try:
            if path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_word_tables(path)
            elif path.suffix.lower() == '.pdf':
                return self._extract_pdf_tables(path, page_number)
            else:
                return {"success": False, "error": f"Unsupported format: {path.suffix}"}
        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            return {"success": False, "error": str(e)}
    
    def get_text_by_section(
        self, 
        file_path: str, 
        start_keyword: Optional[str] = None,
        end_keyword: Optional[str] = None,
        section_title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract text from a specific section of the document.
        
        This is an atomic method that allows the agent to extract specific
        sections (e.g., "Signature", "Summary", "Findings") by keywords.
        
        Args:
            file_path: Path to document
            start_keyword: Keyword that marks the start of the section
            end_keyword: Keyword that marks the end of the section
            section_title: Title of section to find (e.g., "Signature", "Summary")
            
        Returns:
            Dict containing:
            - success: Boolean
            - section_text: Extracted section text
            - start_position: Character position where section starts
            - end_position: Character position where section ends
        """
        # First get full text
        text_result = self.get_full_text(file_path)
        if not text_result.get("success"):
            return text_result
        
        text = text_result["text"]
        
        # If section_title provided, try to find it
        if section_title:
            # Look for section title patterns
            patterns = [
                rf"^{re.escape(section_title)}[:\s]*\n",
                rf"^{re.escape(section_title)}\s*$",
                rf"\n{re.escape(section_title)}[:\s]*\n",
            ]
            
            start_pos = None
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    start_pos = match.end()
                    break
            
            if start_pos is None:
                # Try fuzzy search
                lines = text.split('\n')
                for i, line in enumerate(lines):
                    if section_title.lower() in line.lower():
                        # Get text from this line to next major section
                        start_pos = sum(len(l) + 1 for l in lines[:i])
                        break
            
            if start_pos is not None:
                # Find end of section (next major heading or end of document)
                remaining = text[start_pos:]
                end_match = re.search(r'\n\s*[A-Z][A-Z\s]{10,}\s*\n', remaining)
                end_pos = start_pos + (end_match.start() if end_match else len(remaining))
                
                return {
                    "success": True,
                    "section_text": text[start_pos:end_pos].strip(),
                    "start_position": start_pos,
                    "end_position": end_pos,
                    "section_title": section_title
                }
        
        # Use start/end keywords if provided
        if start_keyword:
            start_match = re.search(re.escape(start_keyword), text, re.IGNORECASE)
            if start_match:
                start_pos = start_match.end()
                if end_keyword:
                    end_match = re.search(re.escape(end_keyword), text[start_pos:], re.IGNORECASE)
                    end_pos = start_pos + (end_match.start() if end_match else len(text) - start_pos)
                else:
                    end_pos = len(text)
                
                return {
                    "success": True,
                    "section_text": text[start_pos:end_pos].strip(),
                    "start_position": start_pos,
                    "end_position": end_pos
                }
        
        return {"success": False, "error": "Could not find specified section"}
    
    def search_text(
        self,
        file_path: str,
        pattern: str, 
        case_sensitive: bool = False,
        use_regex: bool = False,
        context_lines: int = 2
    ) -> Dict[str, Any]:
        """
        Search for text patterns in the document.
        
        This is an atomic method that allows the agent to search for specific
        information (e.g., dates, names, numbers) using patterns.
        
        Args:
            file_path: Path to document
            pattern: Text pattern to search for (string or regex)
            case_sensitive: Whether search is case-sensitive
            use_regex: Whether pattern is a regex pattern
            context_lines: Number of lines of context to include around matches
            
        Returns:
            Dict containing:
            - success: Boolean
            - matches_found: Number of matches
            - matches: List of match dictionaries with:
              - text: Matched text
              - position: Character position
              - context: Surrounding text
              - line_number: Line number where match occurs
        """
        text_result = self.get_full_text(file_path)
        if not text_result.get("success"):
            return text_result
        
        text = text_result["text"]
        lines = text.split('\n')
        matches = []
        
        flags = 0 if case_sensitive else re.IGNORECASE
        
        try:
            if use_regex:
                pattern_obj = re.compile(pattern, flags)
            else:
                pattern_obj = re.compile(re.escape(pattern), flags)
            
            for match in pattern_obj.finditer(text):
                line_num = text[:match.start()].count('\n') + 1
                
                # Get context
                start_line = max(0, line_num - context_lines - 1)
                end_line = min(len(lines), line_num + context_lines)
                context = '\n'.join(lines[start_line:end_line])
                
                matches.append({
                    "text": match.group(0),
                    "position": match.start(),
                    "line_number": line_num,
                    "context": context
                })
            
            return {
                "success": True,
                "pattern": pattern,
                "matches_found": len(matches),
                "matches": matches
            }
        except re.error as e:
            return {"success": False, "error": f"Invalid regex pattern: {e}"}
    
    def extract_structured_data(
        self,
        file_path: str,
        data_types: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Extract common structured data types from the document.
        
        This is an atomic method that extracts dates, names, numbers, emails, etc.
        The agent can use this to quickly find structured information.
        
        Args:
            file_path: Path to document
            data_types: List of data types to extract. Options:
                - "dates": Extract dates
                - "names": Extract potential person names
                - "numbers": Extract numeric values
                - "emails": Extract email addresses
                - "coordinates": Extract coordinate values
                - "depths": Extract depth measurements
                - "all": Extract all types
            
        Returns:
            Dict containing:
            - success: Boolean
            - extracted_data: Dict with keys for each data type
        """
        text_result = self.get_full_text(file_path)
        if not text_result.get("success"):
            return text_result
        
        text = text_result["text"]
        
        if data_types is None:
            data_types = ["all"]
        
        extracted = {}
        
        if "all" in data_types or "dates" in data_types:
            extracted["dates"] = self._extract_dates(text)
        
        if "all" in data_types or "names" in data_types:
            extracted["names"] = self._extract_names(text)
        
        if "all" in data_types or "numbers" in data_types:
            extracted["numbers"] = self._extract_numbers(text)
        
        if "all" in data_types or "emails" in data_types:
            extracted["emails"] = self._extract_emails(text)
        
        if "all" in data_types or "coordinates" in data_types:
            extracted["coordinates"] = self._extract_coordinates(text)
        
        if "all" in data_types or "depths" in data_types:
            extracted["depths"] = self._extract_depths(text)
        
        return {
            "success": True,
            "extracted_data": extracted
        }
    
    # ==========================================================================
    # PRIVATE HELPER METHODS
    # ==========================================================================
    
    def _extract_word_text(self, path: Path, preserve_structure: bool) -> Dict[str, Any]:
        """Extract text from Word document."""
        try:
            doc = docx.Document(path)
            
            if preserve_structure:
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                text = '\n'.join(text_parts)
            else:
                text = '\n'.join([para.text for para in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                text += '\n\n[TABLE]\n'
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    text += row_text + '\n'
                text += '[END TABLE]\n'
            
            # Hard safety cap: never return arbitrarily huge strings as tool output.
            # If Word stats/docProps are missing, size checks may not trigger and this can
            # blow up downstream LLM TPM/context.
            max_chars = 200_000
            truncated = False
            if len(text) > max_chars:
                text = text[:max_chars] + "\n\n[TRUNCATED: document text exceeded safe tool output limit]\n"
                truncated = True
            
            return {
                "success": True,
                "text": text,
                "text_length": len(text),
                "page_count": len(doc.paragraphs) // 50 + 1,
                "structure_preserved": preserve_structure,
                "truncated": truncated,
                "max_chars": max_chars,
            }
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_pdf_text(self, path: Path, preserve_structure: bool) -> Dict[str, Any]:
        """Extract text from PDF document."""
        text = ""
        page_count = 0
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        if preserve_structure:
                            text += page_text + "\n\n"
                        else:
                            text += page_text + " "
            
            # Fallback to PyPDF2 if pdfplumber fails or returns empty
            if not text.strip():
                with open(path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n\n"

            # Hard safety cap (same rationale as Word)
            max_chars = 200_000
            truncated = False
            if len(text) > max_chars:
                text = text[:max_chars] + "\n\n[TRUNCATED: document text exceeded safe tool output limit]\n"
                truncated = True
            
            return {
                "success": True,
                "text": text.strip(),
                "text_length": len(text),
                "page_count": page_count,
                "structure_preserved": preserve_structure,
                "truncated": truncated,
                "max_chars": max_chars,
            }
        except Exception as e:
            logger.error(f"Error reading PDF document: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_word_tables(self, path: Path) -> Dict[str, Any]:
        """Extract tables from Word document."""
        try:
            doc = docx.Document(path)
            tables = []
            
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                # Detect headers (first row if it has text in all cells)
                headers = None
                if table_data and all(cell for cell in table_data[0]):
                    headers = table_data[0]
                
                tables.append({
                    "table_index": table_idx,
                    "row_count": len(table_data),
                    "column_count": len(table_data[0]) if table_data else 0,
                    "data": table_data,
                    "headers": headers
                })
            
            return {
                "success": True,
                "table_count": len(tables),
                "tables": tables
            }
        except Exception as e:
            logger.error(f"Error extracting Word tables: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_pdf_tables(self, path: Path, page_number: Optional[int]) -> Dict[str, Any]:
        """Extract tables from PDF document."""
        try:
            tables = []
            with pdfplumber.open(path) as pdf:
                pages_to_process = [page_number - 1] if page_number else range(len(pdf.pages))
                
                for page_idx in pages_to_process:
                    if 0 <= page_idx < len(pdf.pages):
                        page = pdf.pages[page_idx]
                        page_tables = page.extract_tables()
                        
                        for table_idx, table_data in enumerate(page_tables):
                            if table_data:
                                # Clean up table data
                                cleaned_data = []
                                for row in table_data:
                                    cleaned_row = [cell.strip() if cell else "" for cell in row]
                                    if any(cleaned_row):  # Skip empty rows
                                        cleaned_data.append(cleaned_row)
                                
                                if cleaned_data:
                                    # Detect headers
                                    headers = None
                                    if cleaned_data and all(cell for cell in cleaned_data[0]):
                                        headers = cleaned_data[0]
                                    
                                    tables.append({
                                        "page_number": page_idx + 1,
                                        "table_index": table_idx,
                                        "row_count": len(cleaned_data),
                                        "column_count": len(cleaned_data[0]) if cleaned_data else 0,
                                        "data": cleaned_data,
                                        "headers": headers
                                    })
            
            return {
                "success": True,
                "table_count": len(tables),
                "tables": tables
            }
        except Exception as e:
            logger.error(f"Error extracting PDF tables: {e}")
            return {"success": False, "error": str(e)}
    
    def _extract_dates(self, text: str) -> List[Dict[str, Any]]:
        """Extract dates from text."""
        dates = []
        
        # Common date patterns
        patterns = [
            (r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', 'MM/DD/YYYY'),
            (r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4}', 'DD Month YYYY'),
            (r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{2,4}', 'Month DD, YYYY'),
            (r'\d{4}[/-]\d{1,2}[/-]\d{1,2}', 'YYYY/MM/DD'),
        ]
        
        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                dates.append({
                    "text": match.group(0),
                    "position": match.start(),
                    "format": format_type
                })
        
        return dates
    
    def _extract_names(self, text: str) -> List[Dict[str, Any]]:
        """Extract potential person names from text."""
        names = []
        
        # Pattern: Title? FirstName LastName (common name patterns)
        # This is a simple heuristic - agent should verify
        pattern = r'\b(?:Mr|Mrs|Ms|Dr|Prof|Engr|Surveyor)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'
        
        for match in re.finditer(pattern, text):
            names.append({
                "text": match.group(1),
                "full_match": match.group(0),
                "position": match.start()
            })
        
        # Also look for "Name:" patterns
        name_label_pattern = r'(?:Name|Surveyor|Supervisor|Person)[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)'
        for match in re.finditer(name_label_pattern, text, re.IGNORECASE):
            names.append({
                "text": match.group(1),
                "full_match": match.group(0),
                "position": match.start(),
                "source": "labeled"
            })
        
        return names
    
    def _extract_numbers(self, text: str) -> List[Dict[str, Any]]:
        """Extract numeric values from text."""
        numbers = []
        
        # Pattern for numbers (including decimals)
        pattern = r'\b\d+\.?\d*\b'
        
        for match in re.finditer(pattern, text):
            try:
                value = float(match.group(0))
                numbers.append({
                    "value": value,
                    "text": match.group(0),
                    "position": match.start()
                })
            except ValueError:
                continue
        
        return numbers
    
    def _extract_emails(self, text: str) -> List[Dict[str, Any]]:
        """Extract email addresses from text."""
        emails = []
        pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        
        for match in re.finditer(pattern, text):
            emails.append({
                "text": match.group(0),
                "position": match.start()
            })
        
        return emails
    
    def _extract_coordinates(self, text: str) -> List[Dict[str, Any]]:
        """Extract coordinate values from text."""
        coordinates = []
        
        # UTM coordinates
        utm_pattern = r'E[:\s]*(\d+\.?\d*)\s*N[:\s]*(\d+\.?\d*)'
        for match in re.finditer(utm_pattern, text, re.IGNORECASE):
            coordinates.append({
                "type": "UTM",
                "easting": float(match.group(1)),
                "northing": float(match.group(2)),
                "text": match.group(0),
                "position": match.start()
            })
        
        # Lat/Lon
        latlon_pattern = r'(\d+\.?\d*)[°\s]+([NS])\s+(\d+\.?\d*)[°\s]+([EW])'
        for match in re.finditer(latlon_pattern, text, re.IGNORECASE):
            coordinates.append({
                "type": "LatLon",
                "latitude": float(match.group(1)),
                "longitude": float(match.group(3)),
                "text": match.group(0),
                "position": match.start()
            })
        
        return coordinates
    
    def _extract_depths(self, text: str) -> List[Dict[str, Any]]:
        """Extract depth measurements from text."""
        depths = []
        
        # Pattern: "depth: X m" or "X m deep" or "at X m"
        patterns = [
            r'depth[:\s]+(\d+\.?\d*)\s*(?:m|meters?|ft|feet)',
            r'(\d+\.?\d*)\s*(?:m|meters?|ft|feet)\s+deep',
            r'at\s+(\d+\.?\d*)\s*(?:m|meters?|ft|feet)',
            r'shallowest[:\s]+(\d+\.?\d*)\s*(?:m|meters?|ft|feet)',
        ]
        
        for pattern in patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                try:
                    value = float(match.group(1))
                    depths.append({
                        "value": value,
                        "unit": match.group(2) if len(match.groups()) > 1 else "m",
                        "text": match.group(0),
                        "position": match.start()
                    })
                except (ValueError, IndexError):
                    continue
        
        return depths
    
    # ==========================================================================
    # DOCUMENT STRUCTURE ANALYSIS
    # ==========================================================================
    
    def get_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze document structure: headings, sections, and organization.
        
        This helps the agent understand document organization before processing,
        enabling intelligent section extraction for large documents.
        
        Args:
            file_path: Path to document
            
        Returns:
            Dict containing:
            - success: Boolean
            - sections: List of section dictionaries with:
              - heading: Section heading text
              - level: Heading level (1-9)
              - start_paragraph: Paragraph index where section starts
              - estimated_length: Estimated words in section
            - total_sections: Number of sections found
            - document_outline: Hierarchical structure
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}
        
        try:
            if path.suffix.lower() in ['.docx', '.doc']:
                return self._analyze_word_structure(path)
            elif path.suffix.lower() == '.pdf':
                return self._analyze_pdf_structure(path)
            else:
                return {"success": False, "error": f"Unsupported format: {path.suffix}"}
        except Exception as e:
            logger.error(f"Error analyzing document structure: {e}")
            return {"success": False, "error": str(e)}
    
    def _analyze_word_structure(self, path: Path) -> Dict[str, Any]:
        """Analyze Word document structure."""
        try:
            doc = docx.Document(path)
            sections = []
            current_section = None
            
            for i, para in enumerate(doc.paragraphs):
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # Extract heading level
                    try:
                        level = int(para.style.name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                    
                    # Save previous section if exists
                    if current_section:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        "heading": para.text.strip(),
                        "level": level,
                        "start_paragraph": i,
                        "paragraphs": [],
                        "word_count": 0
                    }
                elif current_section:
                    # Add paragraph to current section
                    current_section["paragraphs"].append(para.text)
                    current_section["word_count"] += len(para.text.split())
            
            # Add last section
            if current_section:
                sections.append(current_section)
            
            # Build outline
            outline = []
            for section in sections:
                indent = "  " * (section["level"] - 1)
                outline.append(f"{indent}• {section['heading']} ({section['word_count']} words)")
            
            return {
                "success": True,
                "sections": sections,
                "total_sections": len(sections),
                "document_outline": "\n".join(outline),
                "has_structure": len(sections) > 0
            }
        except Exception as e:
            logger.error(f"Error analyzing Word structure: {e}")
            return {"success": False, "error": str(e)}
    
    def _analyze_pdf_structure(self, path: Path) -> Dict[str, Any]:
        """Analyze PDF document structure (basic implementation)."""
        # PDF structure analysis is more complex - basic implementation
        try:
            with pdfplumber.open(path) as pdf:
                return {
                    "success": True,
                    "sections": [],
                    "total_sections": 0,
                    "document_outline": f"PDF with {len(pdf.pages)} pages",
                    "has_structure": False,
                    "page_count": len(pdf.pages)
                }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def extract_sections_by_keywords(
        self,
        file_path: str,
        keywords: List[str],
        context_lines: int = 5
    ) -> Dict[str, Any]:
        """
        Extract document sections that match given keywords.
        
        This is useful for large documents where you only need specific sections.
        The method finds sections containing the keywords and extracts them with context.
        
        Args:
            file_path: Path to document
            keywords: List of keywords to search for (case-insensitive)
            context_lines: Number of lines of context around matches
            
        Returns:
            Dict containing:
            - success: Boolean
            - matched_sections: List of section dictionaries
            - total_matches: Number of sections found
            - extracted_text: Combined text from matched sections
        """
        path = Path(file_path)
        if not path.exists():
            return {"success": False, "error": f"File not found: {file_path}"}

        # Word documents: prefer fast XML scan to avoid python-docx slowness on huge table-heavy docs.
        if path.suffix.lower() in [".docx", ".doc"]:
            def iter_lines():
                # XML scan for .docx
                if path.suffix.lower() == ".docx":
                    for t in self._iter_docx_lines_from_xml(str(path)):
                        yield t
                    return

                # Legacy .doc fallback (best effort): use python-docx
                doc = docx.Document(path)
                for para in doc.paragraphs:
                    t = (para.text or "").strip()
                    if t:
                        yield t
                for table in doc.tables:
                    yield "[TABLE]"
                    for row in table.rows:
                        row_text = " | ".join((cell.text or "").strip() for cell in row.cells).strip()
                        if row_text:
                            yield row_text
                    yield "[END TABLE]"

            keywords_lower = [kw.lower() for kw in keywords]
            ctx = max(0, int(context_lines))
            prev = deque(maxlen=ctx)
            lookahead: deque[str] = deque()

            it = iter_lines()
            for _ in range(ctx):
                try:
                    lookahead.append(next(it))
                except StopIteration:
                    break

            matched_sections: list[dict[str, Any]] = []
            extracted_text_parts: list[str] = []
            seen_windows: set[tuple[int, int]] = set()
            line_no = 0

            while lookahead:
                current = lookahead.popleft()
                line_no += 1

                # Maintain lookahead window
                try:
                    lookahead.append(next(it))
                except StopIteration:
                    pass

                current_lower = current.lower()
                if any(kw in current_lower for kw in keywords_lower):
                    context_before = list(prev)
                    context_after = list(lookahead)
                    start_line = max(1, line_no - len(context_before))
                    end_line = line_no + len(context_after)
                    window_key = (start_line, end_line)

                    if window_key not in seen_windows:
                        seen_windows.add(window_key)
                        section_text = "\n".join(context_before + [current] + context_after)
                        matched_sections.append({
                            "line_number": line_no,
                            "context_start": start_line,
                            "context_end": end_line,
                            "text": section_text,
                            "matched_line": current
                        })
                        extracted_text_parts.append(section_text)

                prev.append(current)

            if not matched_sections:
                return {
                    "success": True,
                    "matched_sections": [],
                    "total_matches": 0,
                    "extracted_text": "",
                    "message": f"No sections found matching keywords: {keywords}"
                }

            return {
                "success": True,
                "matched_sections": matched_sections,
                "total_matches": len(matched_sections),
                "extracted_text": "\n\n---\n\n".join(extracted_text_parts),
                "keywords_searched": keywords
            }

        # Non-Word fallback: use full text (may still block for huge docs)
        text_result = self.get_full_text(file_path)
        if not text_result.get("success"):
            return text_result

        lines = text_result["text"].split("\n")
        matched_line_indices = []
        keywords_lower = [kw.lower() for kw in keywords]

        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in keywords_lower):
                matched_line_indices.append(i)

        if not matched_line_indices:
            return {
                "success": True,
                "matched_sections": [],
                "total_matches": 0,
                "extracted_text": "",
                "message": f"No sections found matching keywords: {keywords}"
            }

        matched_sections = []
        extracted_text_parts = []

        for line_idx in matched_line_indices:
            start = max(0, line_idx - context_lines)
            end = min(len(lines), line_idx + context_lines + 1)

            section_text = "\n".join(lines[start:end])
            matched_sections.append({
                "line_number": line_idx + 1,
                "context_start": start + 1,
                "context_end": end,
                "text": section_text,
                "matched_line": lines[line_idx]
            })
            extracted_text_parts.append(section_text)

        return {
            "success": True,
            "matched_sections": matched_sections,
            "total_matches": len(matched_sections),
            "extracted_text": "\n\n---\n\n".join(extracted_text_parts),
            "keywords_searched": keywords
        }
    
    def get_resource_estimation(
        self,
        file_path: str,
        model_name: str = "gpt-5-mini"
    ) -> Dict[str, Any]:
        """
        Estimate resource requirements for processing a document.
        
        This provides warnings and cost estimates before processing large documents.
        
        Args:
            file_path: Path to document
            model_name: LLM model name for cost estimation
            
        Returns:
            Dict containing:
            - success: Boolean
            - file_size_mb: File size in MB
            - word_count: Estimated word count
            - estimated_tokens: Estimated token count
            - estimated_cost: Cost estimation dict
            - warnings: List of warning messages
            - recommendations: List of recommendations
        """
        from utils.cost_estimator import estimate_document_processing_cost_from_tokens
        
        metadata = self.get_document_metadata(file_path)
        if not metadata.get("success"):
            return metadata
        
        path = Path(file_path)
        file_size_mb = path.stat().st_size / (1024 * 1024)
        
        # Token estimate: rely on metadata (Word docProps stats are preferred)
        estimated_tokens = int(metadata.get("estimated_tokens", 0) or 0)
        
        # Generate warnings
        warnings = []
        recommendations = []
        
        if file_size_mb > 10:
            warnings.append(f"⚠️ Large file size: {file_size_mb:.2f} MB")
            recommendations.append("Consider using section extraction instead of full document processing")
        
        if estimated_tokens > 100000:
            warnings.append(f"⚠️ High token count: ~{estimated_tokens:,} tokens")
            recommendations.append("Document may exceed context window. Use chunking or section extraction.")
        
        if estimated_tokens > 200000:
            warnings.append(f"⚠️ Very high token count: ~{estimated_tokens:,} tokens")
            recommendations.append("STRONGLY RECOMMENDED: Extract only relevant sections using document_get_section or extract_sections_by_keywords")
        
        # Estimate cost (token-based; avoids allocating huge strings)
        cost_estimation = estimate_document_processing_cost_from_tokens(
            model_name=model_name,
            document_tokens=estimated_tokens,
            query_tokens=500,
            estimated_iterations=3,
            estimated_output_per_iteration=2000
        )
        
        return {
            "success": True,
            "file_path": str(path.resolve()),
            "file_name": path.name,
            "file_size_mb": round(file_size_mb, 2),
            "file_size_bytes": path.stat().st_size,
            "word_count": metadata.get("word_count", 0),
            "page_count": metadata.get("page_count", 0),
            "estimated_tokens": estimated_tokens,
            "estimated_cost": cost_estimation,
            "warnings": warnings,
            "recommendations": recommendations,
            "model": model_name
        }
    
    # ==========================================================================
    # BACKWARD COMPATIBILITY METHODS (for existing code)
    # ==========================================================================
    
    def read_word(self, file_path: str) -> str:
        """Backward compatibility: Read Word document text."""
        result = self.get_full_text(file_path, preserve_structure=False)
        if result.get("success"):
            return result["text"]
        raise FileNotFoundError(result.get("error", "Unknown error"))
    
    def read_pdf(self, file_path: str) -> str:
        """Backward compatibility: Read PDF document text."""
        result = self.get_full_text(file_path, preserve_structure=False)
        if result.get("success"):
            return result["text"]
        raise FileNotFoundError(result.get("error", "Unknown error"))
    
    def process_word(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Backward compatibility: Process Word document."""
        result = self.get_full_text(file_path, preserve_structure=True)
        if not result.get("success"):
            return result
        
        return {
            "file_path": file_path,
            "file_type": "Word",
            "text_length": result["text_length"],
            "text": result["text"],
            "extracted_data": {}
        }
    
    def process_pdf(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Backward compatibility: Process PDF document."""
        result = self.get_full_text(file_path, preserve_structure=True)
        if not result.get("success"):
            return result
        
        return {
            "file_path": file_path,
            "file_type": "PDF",
            "text_length": result["text_length"],
            "text": result["text"],
            "extracted_data": {}
        }
    
    # ==========================================================================
    # DOCUMENT CREATION METHODS
    # ==========================================================================
    
    def create_word_document(
        self,
        file_path: str,
        content: str,
        title: Optional[str] = None,
        add_table_of_contents: bool = False
    ) -> Dict[str, Any]:
        """
        Create a new Word document with the specified content.
        
        This is an atomic method that allows the agent to create Word documents
        programmatically. Useful for generating reports, summaries, or exports.
        
        Args:
            file_path: Path where the document should be saved (.docx)
            content: Text content to write to the document
            title: Optional title for the document
            add_table_of_contents: Whether to add a table of contents
            
        Returns:
            Dict containing:
            - success: Boolean
            - file_path: Full path to created file
            - file_size: Size of created file in bytes
            - message: Success message
        """
        try:
            path = Path(file_path)
            
            # Ensure .docx extension
            if path.suffix.lower() != '.docx':
                path = path.with_suffix('.docx')
            
            # Create parent directories if needed
            path.parent.mkdir(parents=True, exist_ok=True)
            
            # Create new document
            doc = docx.Document()
            
            # Add title if provided
            if title:
                title_para = doc.add_heading(title, 0)
                title_para.alignment = 1  # Center alignment
            
            # Process content - split by lines and preserve structure
            lines = content.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines but add spacing
                if not line:
                    if current_paragraph:
                        doc.add_paragraph()  # Add blank line
                    continue
                
                # Detect headings (lines that are all caps or start with #)
                if line.startswith('#'):
                    # Markdown-style heading
                    level = len(line) - len(line.lstrip('#'))
                    heading_text = line.lstrip('#').strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=min(level, 9))
                elif line.isupper() and len(line) > 5 and not line.startswith('HTTP'):
                    # All caps line might be a heading
                    doc.add_heading(line, level=2)
                else:
                    # Regular paragraph
                    para = doc.add_paragraph(line)
                    current_paragraph = para
            
            # Save document
            doc.save(str(path))
            
            file_size = path.stat().st_size
            
            logger.info(f"Created Word document: {path} ({file_size} bytes)")
            
            return {
                "success": True,
                "file_path": str(path.resolve()),
                "file_size": file_size,
                "message": f"Successfully created Word document: {path.name}"
            }
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to create Word document: {e}"
            }
    
    def read_existing_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Read an existing Word document to get its content.
        
        This allows the agent to read a document it created earlier and update it.
        
        Args:
            file_path: Path to existing Word document
            
        Returns:
            Dict containing:
            - success: Boolean
            - text: Full text content
            - paragraphs: List of paragraph texts
            - tables: List of table data
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return {"success": False, "error": f"File not found: {file_path}"}
            
            doc = docx.Document(path)
            
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = '\n'.join(paragraphs)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            
            return {
                "success": True,
                "text": text,
                "paragraphs": paragraphs,
                "tables": tables,
                "file_path": str(path.resolve())
            }
        except Exception as e:
            logger.error(f"Error reading Word document: {e}")
            return {"success": False, "error": str(e)}
    
    def update_word_document(
        self,
        file_path: str,
        new_content: str,
        title: Optional[str] = None,
        overwrite: bool = True
    ) -> Dict[str, Any]:
        """
        Update an existing Word document with new content.
        
        Args:
            file_path: Path to document (will be created if doesn't exist)
            new_content: New content to write
            title: Optional title
            overwrite: If True, replace entire document; if False, append
            
        Returns:
            Dict with success status and file info
        """
        if overwrite or not Path(file_path).exists():
            # Create new document
            return self.create_word_document(file_path, new_content, title)
        else:
            # Append to existing
            existing = self.read_existing_word_document(file_path)
            if existing.get("success"):
                combined_content = existing["text"] + "\n\n" + new_content
                return self.create_word_document(file_path, combined_content, title)
            else:
                return existing
    
    def create_word_document_from_structure(
        self,
        file_path: str,
        title: str,
        sections: List[Dict[str, Any]],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Create a Word document from structured data.
        
        This is an atomic method for creating professional documents with
        structured sections, tables, and formatting.
        
        Args:
            file_path: Path where document should be saved
            title: Document title
            sections: List of section dictionaries, each with:
                - heading: Section heading text
                - level: Heading level (1-9)
                - content: Section content (text or list of paragraphs)
                - table: Optional table data (list of lists)
            metadata: Optional metadata dict (author, date, etc.)
            
        Returns:
            Dict with success status and file info
        """
        try:
            path = Path(file_path)
            if path.suffix.lower() != '.docx':
                path = path.with_suffix('.docx')
            
            path.parent.mkdir(parents=True, exist_ok=True)
            
            doc = docx.Document()
            
            # Add title
            doc.add_heading(title, 0)
            
            # Add metadata if provided
            if metadata:
                meta_para = doc.add_paragraph()
                for key, value in metadata.items():
                    meta_para.add_run(f"{key}: {value}\n")
                doc.add_paragraph()  # Blank line
            
            # Add sections
            for section in sections:
                heading = section.get("heading", "")
                level = section.get("level", 2)
                content = section.get("content", "")
                table_data = section.get("table")
                
                if heading:
                    doc.add_heading(heading, level=level)
                
                if content:
                    if isinstance(content, list):
                        for para_text in content:
                            if para_text:
                                doc.add_paragraph(para_text)
                    else:
                        doc.add_paragraph(content)
                
                if table_data:
                    # Create table
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 0)
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = str(cell_data)
                
                doc.add_paragraph()  # Blank line between sections
            
            doc.save(str(path))
            file_size = path.stat().st_size
            
            logger.info(f"Created structured Word document: {path} ({file_size} bytes)")
            
            return {
                "success": True,
                "file_path": str(path.resolve()),
                "file_size": file_size,
                "message": f"Successfully created Word document: {path.name}"
            }
            
        except Exception as e:
            logger.error(f"Error creating structured Word document: {e}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to create Word document: {e}"
            }
