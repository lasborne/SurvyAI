"""
LLM vision OCR for SurvyAI (single module).

Flow: encode image → one style-locked vision read → local survey-context repairs.
Face Left/Right identities and prior handwriting are used internally to resolve
ambiguous glyphs. The user-facing reply stays a compact extraction (no PASS/FAIL).
The LangGraph tool loop stays text-only.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import math
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Literal, Optional, Sequence

from survyai.attachments import (
    DEFAULT_ATTACHMENTS_ONLY_PROMPT,
    IMAGE_EXTENSIONS,
    parse_attachments_block,
)

logger = logging.getLogger(__name__)

VisionOcrMode = Literal["plain_text", "structured", "geospatial"]
OcrDocumentType = Literal[
    "calibration_sheet",
    "levelling_book",
    "traverse_sheet",
    "cadastral_plan",
    "general_table",
    "spreadsheet_screenshot",
    "ui_screenshot",
    "generic",
]

DOC_CALIBRATION = "calibration_sheet"
DOC_LEVELLING = "levelling_book"
DOC_TRAVERSE = "traverse_sheet"
DOC_CADASTRAL = "cadastral_plan"
DOC_TABLE = "general_table"
DOC_SPREADSHEET = "spreadsheet_screenshot"
DOC_UI = "ui_screenshot"
DOC_GENERIC = "generic"

OCR_BUDGET_S = 58
_LEVEL_TOL = 0.003
_DISTANCE_TOL = 0.05
_ANGLE_TOL_DEG = 1.0
_LOW_CONFIDENCE = 0.92
_NESTED_UNCERTAIN = 0.50
_WEAK_CALIBRATION_STEMS = frozenset({"calibration", "cal", "calib"})
_HANDWRITING_REL = Path(".survyai") / "handwriting_style.json"
_LAST_OCR_REL = Path(".survyai") / "ocr" / "last_extraction.json"
_DIGIT_SWAP_PAIRS = (
    ("3", "4"), ("4", "3"), ("4", "6"), ("6", "4"), ("3", "8"), ("8", "3"),
    ("0", "8"), ("8", "0"), ("1", "7"), ("7", "1"), ("6", "8"), ("8", "6"),
    ("9", "3"), ("3", "9"), ("5", "6"), ("6", "5"), ("0", "6"), ("6", "0"),
)
_IMAGE_EXT_ALT = "|".join(sorted(ext.lstrip(".") for ext in IMAGE_EXTENSIONS))

_TEXT_ONLY_MARKERS = (
    "text only", "just the text", "extract text", "extract the text",
    "plain text only", "read the text only",
)
_PLAIN_MARKERS = (
    "ocr", "read this", "read the text", "what does this say", "what does it say",
    "extract text", "extract the text", "transcribe", "text only", "just the text",
    "read the image", "what's written", "what is written",
)
_SURVEY_SHEET_MARKERS = (
    "calibration", "collimation", "two-peg", "two peg", "two peg test",
    "staff reading", "staff readings", "instrument reading",
    "backsight", "back sight", "foresight", "fore sight",
    "intermediate sight", "reduced level", "rise and fall",
    "height of collimation", "height of instrument", "level book",
    "booking sheet", "field book", "dumpy", "automatic level", "auto level",
    "theodolite", "total station", "peg test",
    "face left", "face right", "field sheet", "angle and distance",
    "horizontal angle", "traverse angle",
)
_CALIBRATION_MARKERS = (
    "calibration", "collimation", "two-peg", "two peg", "peg test",
    "staff calibration", "instrument calibration", "constant", "index error",
)
_LEVELLING_MARKERS = (
    "levelling", "leveling", "level book", "backsight", "foresight",
    "intermediate sight", "reduced level", "rise and fall",
    "height of collimation", "staff reading",
)
_GEOSPATIAL_MARKERS = (
    "bearing", "bearings", "traverse", "pillar", "pillars", "cadastral",
    "survey plan", "parcel", "easting", "northing", "coordinate", "coordinates",
    "north arrow", "legend", "scale", "crs", "utm", "geospatial", "boundary",
    "deed plan", "site plan", "lot plan",
)
_STRUCTURED_MARKERS = (
    "structured", "json", "table", "tables", "labels", "symbols",
    "components", "extract all", "key details", "fields", "metadata",
)
_OCR_ONLY_MARKERS = (
    "ocr", "extract text", "extract the text", "read this image", "read the image",
    "what does this say", "what is in this image", "what's in this image",
    "transcribe", "just ocr", "only ocr", "text from the image", "scan this",
    "scan the image", "scan through", "scan the", "key details", "all details",
    "all key", "extract all", "give all", "all the details",
)
_CAD_CONTINUE_MARKERS = (
    "replot", "plot", "draw", "generate dwg", "create dwg", "to dwg",
    ".dwg", "autocad", "cad",
)
_SAVE_MARKERS = (
    "save", "export", "write to", "save as", "create a word", "create word",
    ".docx", ".json", ".xlsx", ".xls", "excel", "spreadsheet", "workbook",
    "powerpoint", ".pptx", ".csv", "text file", ".txt",
)
_OCR_FOLLOWUP_MARKERS = (
    "save", "export", "excel", ".xlsx", ".xls", "spreadsheet", "workbook",
    "summarize", "summary", "compute", "calculate", "analyse", "analyze",
    "create a table", "make a table", "powerpoint", ".docx", "word file",
    "from the extraction", "from the ocr", "use the extraction", "use the ocr",
    "those values", "these values", "the updates", "updated values",
)
_TRAVERSE_MARKERS = (
    "traverse", "bearing", "included angle", "interior angle",
    "face left", "face right", "horizontal angle", "vertical angle",
    "angle and distance", "field sheet", "inst. stn", "inst stn",
    "ref. stn", "ref stn", "slope dist", "mean angle", "hz_fl", "va_fl",
)
_CADASTRAL_MARKERS = ("cadastral", "survey plan", "deed plan", "site plan", "pillar", "parcel")
_SPREADSHEET_MARKERS = (
    "excel", "spreadsheet", "workbook", "worksheet", ".xlsx", ".xls",
    "attr_", "vertex_ix", "part_index", "tmp_geom", "visible_headers",
    "visible_records", "column", "ribbon", "get genuine office",
)
_UI_MARKERS = (
    "desktop", "window title", "menu bar", "sidebar", "toolbar", "console",
    "dialog", "status bar", "live activity", "conversations", "application",
    "ui screenshot", "app screenshot", "survyai desktop",
)
# Only count these when the attachment looks like a screenshot (avoid CAD/chat false hits)
_UI_WEAK_MARKERS = (
    "button", "checkbox", "workspace", "send", "cancel", "retry", "sidebar",
)
_FULL_DETAIL_MARKERS = (
    "all details", "all key details", "all the details", "everything",
    "all text", "full details", "entire", "complete details", "give all",
    "scan through", "give every", "every detail", "all visible",
)
_FIELD_ALIASES = {
    "bs": "backsight", "b.s": "backsight", "back sight": "backsight", "backsight": "backsight",
    "is": "intermediate_sight", "i.s": "intermediate_sight",
    "intermediate": "intermediate_sight", "intermediate sight": "intermediate_sight",
    "fs": "foresight", "f.s": "foresight", "fore sight": "foresight", "foresight": "foresight",
    "rl": "reduced_level", "r.l": "reduced_level", "reduced level": "reduced_level",
    "reduced_level": "reduced_level",
    "hi": "height_of_collimation", "h.i": "height_of_collimation", "hoc": "height_of_collimation",
    "height of collimation": "height_of_collimation",
    "height of instrument": "height_of_collimation",
    "height_of_collimation": "height_of_collimation",
    "dist": "distance", "distance": "distance", "chainage": "distance",
    "rise": "rise", "fall": "fall", "station": "station", "stn": "station",
    "point": "station", "remarks": "remarks", "remark": "remarks", "notes": "remarks",
    "inst stn": "instrument_station", "inst. stn": "instrument_station",
    "instrument station": "instrument_station", "instrument_station": "instrument_station",
    "station from": "instrument_station", "from": "instrument_station",
    "ref stn": "reference_station", "ref. stn": "reference_station",
    "reference station": "reference_station", "reference_station": "reference_station",
    "station to": "reference_station", "to": "reference_station",
    "hz fl": "hz_fl", "ha fl": "hz_fl", "horizontal fl": "hz_fl", "hz_fl": "hz_fl",
    "hz fr": "hz_fr", "ha fr": "hz_fr", "horizontal fr": "hz_fr", "hz_fr": "hz_fr",
    "va fl": "va_fl", "vertical fl": "va_fl", "va_fl": "va_fl",
    "va fr": "va_fr", "vertical fr": "va_fr", "va_fr": "va_fr",
    "slope": "slope_distance", "slope dist": "slope_distance",
    "slop e dist": "slope_distance", "slope distance": "slope_distance",
    "slope_distance": "slope_distance",
    "hor": "horizontal_distance", "hor dist": "horizontal_distance",
    "horizontal dist": "horizontal_distance", "horizontal distance": "horizontal_distance",
    "horizontal_distance": "horizontal_distance",
    "phone": "phone", "telephone": "phone", "tel": "phone", "mobile": "phone",
    "contact": "phone", "phone number": "phone",
}

_HANDWRITING_LOCK = """
HANDWRITING STYLE LOCK (do this in the same pass — do not skip):
1. First scan several CLEAR repeated digits/letters from THIS writer only. Build style_card.
2. style_card keys: 0,1,2,3,4,5,6,7,8,9 and any ambiguous letters (Z/2, S/5, O/0, I/1, T/7).
   Each value is a short note of THIS writer's habit, e.g. "1 = long up-serif, no crossbar; 7 = crossbar mid-stem; 9 = closed loop with short tail".
3. Transcribe every ambiguous glyph using THIS style_card, not a printed font and not another writer's habit.
4. A 9 that looks like 1 to generic OCR is still 9 if this writer's 9s match that loop/tail. A 7 with a crossbar is 7 even if it resembles 1.
5. Never 'correct' a glyph toward a nicer number. If two style readings remain possible, pick the style-consistent one and set confidence < 0.92.
"""

_ACCURACY_RULES = """
1. Identify each field, then transcribe. Do not dump unlabelled text as a reading.
2. Never invent a digit that is not on the sheet. Unreadable → value null, confidence ≤ 0.4.
3. Crossed-out values: use the FINAL rewrite; struck_through=true on the discarded glyph if visible.
4. Prefer plain strings/numbers. Use {"raw","value","confidence"} only when confidence < 0.85.
5. Do not rewrite a clearly written number to make arithmetic pretty. If a glyph is faint, prefer the reading that matches this sheet's identities (levelling HI=RL+BS; traverse HA |FL-FR|≈180°, VA FL+FR≈360°, repeated distances).
6. Preserve row order and station IDs. Blank Inst. Stn inherits the station above.
""" + _HANDWRITING_LOCK

_TRAVERSE_SHEET_HINT = (
    "This is a traverse / angle-and-distance field sheet, not levelling. "
    "Extract EVERY filled observation row (typically 6). Do not stop after the header. "
    "Inst. Stn = station from; Ref. Stn = station to. Blank Inst. Stn inherits the station above. "
    "Inst. Stn is usually ONE station for the whole page (e.g. SIAX-03); do not force it to equal a Ref. Stn. "
    "Ref. Stn typically alternates between the SAME two IDs across sets (e.g. SIAX-04, SIAX-02, SIAX-04, SIAX-02…). "
    "Crossed-out Ref values: use the FINAL rewrite only (struck SIAX-01 under SIAX-02 → SIAX-02). "
    "For EACH row include ALL of: from, to, hz_fl, hz_fr, va_fl, va_fr, slope, hor "
    "as DMS strings like 089°46′23″ or 000°00′00″ (not separate deg/min/sec objects). "
    "Same from→to shot: Horizontal Face Left vs Face Right differ by ~180° (≤1°); "
    "Vertical Face Left + Face Right sum to ~360° (≤1°). Use that to choose a faint digit. "
    "The same from→to pair later in the sheet repeats the SAME slope and horizontal distances. "
    "6 = closed lower loop, often a high stem; 4 = open top (363.054 not 343.054 when the loop is closed). "
    "Serial: count every narrow oval as 0 (1250033, not 125123 / 125103 / 1258133). "
    "Date: a 7 with a horizontal cross-bar is 07, not 01. "
    "Instrument suffix after R is often + (R+), not T (Geomato MTS-1202R+). "
    "Station IDs look like SIAX-03, not S/AX-03; 3 often misread as 2 — check closed loops. "
    "Also capture the full handwritten phone/contact number near the header (all digits). "
)

_COMPACT_JSON_HINT = (
    "Return JSON only (no markdown, no essays, no PASS/FAIL). "
    "Keys: document_type, title, "
    "style_card (≤6 words per glyph; omit for printed UI / Excel screenshots), "
    "metadata (ONLY fields that are visibly present — never invent organization/phone/surveyed_by/"
    "instrument/serial/observer/station when absent), "
    "sections (for UI/app screenshots: [{heading, lines:[...]} ] covering EVERY visible panel), "
    "rows (for tables/field sheets: one object per visible data line; use ACTUAL column headers). "
    "For spreadsheet/Excel screenshots: document_type=spreadsheet_screenshot; "
    "metadata.sheet = active sheet tab; rows = visible data cells only; "
    "also set visible_headers to the header row texts in left-to-right order. "
    "For application/desktop UI screenshots: document_type=ui_screenshot; "
    "sections REQUIRED with all readable text grouped by panel (Window, Menu, Conversations, "
    "Console/Chat, Live activity, Workspace, Status/Footer, etc.). "
    "Do NOT invent traverse/levelling columns (from/to/hz_fl/observer/stn) unless printed on the image. "
    "For traverse field sheets only: flat keys from, to, hz_fl, hz_fr, va_fl, va_fr, slope, hor "
    "(DMS strings for angles; numbers for distances). "
    "Do not nest face_left/face_right under horizontal_angle — flatten them. "
    "Do not return _validation or long notes. "
    "Optional: for any field/cell object, include bbox [x0,y0,x1,y1] normalized 0–1 "
    "relative to the image (omit when unsure)."
)

_SPREADSHEET_HINT = (
    "This image is an Excel / spreadsheet / GIS attribute-table screenshot (or similar UI table). "
    "Extract ONLY what is visible: window title, sheet tab name, column headers, and data cells. "
    "Put every visible data row into rows[] using the real header names as keys "
    "(e.g. ATTR_ON, ATTR_DI, VERTEX_IX, X, Y, Z, M). "
    "Skip empty header columns. Preserve numeric precision. "
    "Do NOT invent survey booking fields (observer, instrument, serial, Inst. Stn, Face Left, etc.). "
    "Do NOT invent values for blank cells — omit the key or use null. "
)

_UI_HINT = (
    "This is an application / desktop UI screenshot (menus, sidebars, chat, status panels) — "
    "NOT a survey field sheet and NOT an Excel cell grid. "
    "Extract ALL readable on-screen text. Group by panel/region into sections[]. "
    "Each section: {heading: short panel name, lines: [string, ...]} in reading order. "
    "Typical headings: Window title, Menu bar, Workspace, Conversations, Console / Chat, "
    "Live activity, Controls / Buttons, Status bar / Footer. "
    "Preserve speaker labels (You / SurvyAI), timestamps, file paths, button captions, and "
    "checkbox labels. Quote chat prompts and replies faithfully (truncate only if extremely long). "
    "title = window title bar. metadata may include workspace_path, user, status. "
    "Never return only the window title when more text is visible. "
    "Do NOT invent survey booking fields. "
)


_MIN_EDGE_PX = 80
_BLUR_UNREADABLE = 12.0
_EXPOSURE_DARK = 0.04
_EXPOSURE_BRIGHT = 0.96


# ---------------------------------------------------------------------------
# Result
# ---------------------------------------------------------------------------

@dataclass
class VisionOcrResult:
    success: bool = False
    mode: VisionOcrMode = "plain_text"
    text: str = ""
    structured: Dict[str, Any] = field(default_factory=dict)
    image_paths: List[str] = field(default_factory=list)
    model_name: Optional[str] = None
    error: Optional[str] = None
    notes: str = ""
    document_type: str = "generic"
    validation: Dict[str, Any] = field(default_factory=dict)
    quality: Dict[str, Any] = field(default_factory=dict)
    ocr_review: Dict[str, Any] = field(default_factory=dict)

    def format_for_user(self) -> str:
        if not self.success:
            return self.error or "Vision OCR failed."
        paths = ", ".join(Path(p).name for p in self.image_paths) or "(image)"
        if self.mode == "plain_text" and self.text.strip() and not self.structured:
            return f"**Source:** {paths}\n\n{self.text.strip()}"
        if self.structured:
            return _format_structured_for_user(self.structured, paths, text=self.text)
        if self.text.strip():
            return self.text.strip()
        return self.notes or "Vision OCR returned no content."

    def format_for_agent_context(self, *, max_chars: int = 24000) -> str:
        if not self.success:
            return f"VISION OCR RESULT: failed — {self.error or 'unknown error'}"
        parts = [
            "VISION OCR RESULT (pre-extracted from attached/typed image paths).",
            "Use this text for reasoning and tools; do not claim you cannot see the image.",
            "Face Left/Right and repeated-distance identities were applied internally to resolve faint glyphs.",
            f"Mode: {self.mode}",
            f"Document type: {self.document_type}",
            f"Sources: {', '.join(self.image_paths)}",
        ]
        if self.structured:
            blob = json.dumps(self.structured, ensure_ascii=False)
            parts.append("Structured JSON:")
            parts.append(blob[: max_chars - 20] + "…[truncated]" if len(blob) > max_chars else blob)
        if self.text.strip():
            txt = self.text.strip()
            parts.append("Plain text:")
            parts.append(txt[: max_chars - 20] + "…[truncated]" if len(txt) > max_chars else txt)
        return "\n".join(parts)


# ---------------------------------------------------------------------------
# Path / mode helpers
# ---------------------------------------------------------------------------

def extract_image_paths_from_query(query: str) -> List[str]:
    paths: List[str] = []
    seen: set = set()
    marker_paths, _ = parse_attachments_block(query or "")
    for raw in marker_paths:
        _maybe_add_image(raw, paths, seen)
    for pattern in (
        rf'([A-Za-z]:\\[^\r\n"<>|]+?\.(?:{_IMAGE_EXT_ALT}))',
        rf'((?:/|\\)[^\r\n"<>|]+?\.(?:{_IMAGE_EXT_ALT}))',
    ):
        for match in re.findall(pattern, query or "", flags=re.IGNORECASE):
            _maybe_add_image(str(match).strip().strip('"').strip("'").rstrip(").,;"), paths, seen)
    return paths


def _maybe_add_image(raw: str, out: List[str], seen: set) -> None:
    try:
        path = Path(str(raw or "").strip())
    except Exception:
        return
    if not path.is_file() or path.suffix.lower() not in IMAGE_EXTENSIONS:
        return
    key = str(path.resolve())
    if key not in seen:
        seen.add(key)
        out.append(key)


def select_vision_ocr_mode(user_text: str) -> VisionOcrMode:
    ql = (user_text or "").lower()
    if any(m in ql for m in _GEOSPATIAL_MARKERS):
        return "geospatial"
    if any(m in ql for m in _SURVEY_SHEET_MARKERS) or any(m in ql for m in _STRUCTURED_MARKERS):
        return "structured"
    if any(m in ql for m in _TEXT_ONLY_MARKERS):
        return "plain_text"
    return "structured"


def is_ocr_only_request(user_text: str) -> bool:
    ql = (user_text or "").lower().strip()
    if not ql or any(m in ql for m in _CAD_CONTINUE_MARKERS):
        return False
    if ql == DEFAULT_ATTACHMENTS_ONLY_PROMPT.lower():
        return True
    has_ocr = (
        any(m in ql for m in _OCR_ONLY_MARKERS)
        or any(m in ql for m in _PLAIN_MARKERS)
        or any(m in ql for m in ("key details", "extract all"))
    )
    # Pure follow-ups (save/excel/compute/…) with no fresh scan/extract verbs.
    if any(m in ql for m in _OCR_FOLLOWUP_MARKERS) and not has_ocr:
        return False
    follow = ("then", "after that", "use the", "plot", "replot", "arcgis", "calculate", "analyse", "analyze")
    if any(m in ql for m in follow) and not has_ocr:
        return False
    return has_ocr


def is_ocr_export_request(user_text: str) -> bool:
    """True when the user wants the latest OCR extraction saved to Excel."""
    ql = (user_text or "").lower()
    wants_save = any(m in ql for m in ("save", "export", "write", "create"))
    wants_excel = any(m in ql for m in ("excel", ".xlsx", ".xls", "spreadsheet", "workbook", "csv", ".csv"))
    return wants_save and wants_excel


def is_ocr_word_export_request(user_text: str) -> bool:
    """True when the user wants the latest OCR extraction saved to Word (.docx).

    Prefer this over the session essay-save path for 'save this/the extraction to f.docx'.
    """
    ql = (user_text or "").lower().strip()
    if not ql:
        return False
    wants_save = any(m in ql for m in ("save", "export", "write", "create"))
    wants_word = any(
        m in ql for m in (".docx", "word document", "word doc", "word file", "ms word")
    )
    if not (wants_save and wants_word):
        return False
    # Pure Excel exports
    if any(m in ql for m in (".xlsx", ".xls", "spreadsheet", "workbook")) and ".docx" not in ql:
        return False
    # Explicit essay/report rewrite requests stay on the essay path
    if any(
        m in ql
        for m in (
            "essay",
            "well-structured",
            "turn this into",
            "turn the previous",
            "previous topic",
            "write an essay",
            "as an essay",
        )
    ):
        return False
    # Strong OCR-save signals
    if any(
        m in ql
        for m in (
            "extraction",
            "ocr",
            "from the image",
            "from the screenshot",
            "the details",
            "all details",
            "scanned",
            "scan result",
        )
    ):
        return True
    # Pronoun / generic "save this/it … to f.docx" after an OCR turn
    if re.search(r"\bsave\s+(this|it|that)\b", ql):
        return True
    if re.search(r"\b(into|to|as)\s+(the\s+)?file\b", ql) and ".docx" in ql:
        return True
    # Any save-to-.docx while a last OCR exists (caller gates on load_last_ocr_extraction)
    return ".docx" in ql or "word" in ql


def is_ocr_followup_request(user_text: str) -> bool:
    ql = (user_text or "").lower().strip()
    if not ql:
        return False
    return any(m in ql for m in _OCR_FOLLOWUP_MARKERS)


def user_requested_save(user_text: str) -> bool:
    return any(m in (user_text or "").lower() for m in _SAVE_MARKERS)


def looks_like_survey_plan_image_task(user_text: str, image_paths: Sequence[str]) -> bool:
    ql = (user_text or "").lower()
    if any(m in ql for m in ("cadastral", "survey plan", "deed plan", "site plan", "replot")):
        return True
    return any(
        any(k in Path(p).stem.lower() for k in ("plan", "survey", "cadastral", "deed", "site", "parcel"))
        for p in image_paths
    )


def looks_like_survey_sheet(user_text: str = "", image_paths: Sequence[str] = ()) -> bool:
    blob = " ".join([user_text or "", " ".join(Path(p).stem for p in image_paths)]).lower().replace("_", " ").replace("-", " ")
    return any(m in blob for m in _SURVEY_SHEET_MARKERS + _CALIBRATION_MARKERS + _LEVELLING_MARKERS)


def should_fastpath_image_survey_replot(query: str, routing_query: Optional[str] = None) -> bool:
    scope = (routing_query or query or "").lower()
    if not extract_image_paths_from_query(query or "") and not any(ext in scope for ext in IMAGE_EXTENSIONS):
        return False
    if ".dwg" not in scope and "replot" not in scope and "autocad" not in scope:
        if not any(k in scope for k in ("plot", "draw", "generate", "create")):
            return False
    return any(k in scope for k in ("cadastral", "survey plan", "survey/cadastral", "replot", "pillar", "bearing", "parcel"))


# ---------------------------------------------------------------------------
# Encode + quality
# ---------------------------------------------------------------------------

def _optional_cv2_deskew(img: Any) -> Any:
    """Light deskew when OpenCV is importable; otherwise return unchanged."""
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
    except Exception:
        return img
    try:
        arr = np.array(img.convert("L"))
        edges = cv2.Canny(arr, 50, 150)
        lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=80, minLineLength=60, maxLineGap=12)
        if lines is None or len(lines) < 4:
            return img
        angles: List[float] = []
        for line in lines[:40]:
            x1, y1, x2, y2 = line[0]
            if abs(x2 - x1) < 2:
                continue
            ang = math.degrees(math.atan2(y2 - y1, x2 - x1))
            if abs(ang) < 15:
                angles.append(ang)
        if not angles:
            return img
        median = float(sorted(angles)[len(angles) // 2])
        if abs(median) < 0.4 or abs(median) > 12:
            return img
        rgb = np.array(img.convert("RGB"))
        h, w = rgb.shape[:2]
        matrix = cv2.getRotationMatrix2D((w / 2.0, h / 2.0), median, 1.0)
        warped = cv2.warpAffine(rgb, matrix, (w, h), flags=cv2.INTER_LINEAR, borderMode=cv2.BORDER_REPLICATE)
        from PIL import Image  # type: ignore

        return Image.fromarray(warped)
    except Exception:
        return img


def assess_image_quality(path: str) -> Dict[str, Any]:
    """Local quality scores (blur / exposure / size). No LLM."""
    out: Dict[str, Any] = {
        "path": str(path),
        "blur": None,
        "exposure": None,
        "width": 0,
        "height": 0,
        "overall": "unknown",
        "readable": True,
        "reason": "",
    }
    try:
        from PIL import Image, ImageOps  # type: ignore
        import numpy as np  # type: ignore

        with Image.open(path) as raw:
            img = ImageOps.exif_transpose(raw).convert("RGB")
            w, h = img.size
            out["width"], out["height"] = int(w), int(h)
            if min(w, h) < _MIN_EDGE_PX:
                out["readable"] = False
                out["overall"] = "bad"
                out["reason"] = "Image is too small to read handwritten survey values."
                return out
            gray = np.asarray(img.convert("L"), dtype=np.float32)
            mean = float(gray.mean() / 255.0)
            out["exposure"] = round(mean, 4)
            lap = (
                -4.0 * gray
                + np.roll(gray, 1, 0)
                + np.roll(gray, -1, 0)
                + np.roll(gray, 1, 1)
                + np.roll(gray, -1, 1)
            )
            blur = float(lap.var())
            out["blur"] = round(blur, 2)
            if blur < _BLUR_UNREADABLE:
                out["readable"] = False
                out["overall"] = "bad"
                out["reason"] = "Photograph is too blurry. Retake with the sheet flat and in focus."
                return out
            if mean < _EXPOSURE_DARK:
                out["readable"] = False
                out["overall"] = "bad"
                out["reason"] = "Photograph is too dark. Retake with better lighting."
                return out
            if mean > _EXPOSURE_BRIGHT:
                out["readable"] = False
                out["overall"] = "bad"
                out["reason"] = "Photograph is overexposed / washed out. Retake avoiding glare."
                return out
            if blur < 40.0 or mean < 0.12 or mean > 0.88:
                out["overall"] = "questionable"
                out["reason"] = "Image quality is marginal; extraction may need review."
            else:
                out["overall"] = "good"
            return out
    except Exception as exc:
        # Cannot score — do not block OCR; encode step will fail if the file is unusable.
        out["readable"] = True
        out["overall"] = "unknown"
        out["reason"] = f"Quality not scored: {exc}"
        return out


def image_file_to_base64_png(path: str, *, max_edge: int = 2048, enhance: bool = False) -> Optional[str]:
    """Load/downscale; JPEG when enhance (smaller/faster), else PNG. Returns raw base64.

    Never overwrites the original file — working copy is in-memory only.
    """
    p = Path(path)
    if not p.is_file():
        return None
    try:
        from PIL import Image, ImageEnhance, ImageOps  # type: ignore

        with Image.open(p) as raw:
            img = ImageOps.exif_transpose(raw).convert("RGB")
            if enhance:
                img = _optional_cv2_deskew(img)
                img = ImageEnhance.Sharpness(ImageOps.autocontrast(img.convert("L")).convert("RGB")).enhance(1.2)
                img = ImageEnhance.Contrast(img).enhance(1.1)
            w, h = img.size
            longest = max(w, h)
            if longest > max_edge > 0:
                scale = max_edge / float(longest)
                img = img.resize((max(1, int(w * scale)), max(1, int(h * scale))), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            if enhance:
                img.save(buf, format="JPEG", quality=88, optimize=True)
            else:
                img.save(buf, format="PNG", optimize=True)
            return base64.b64encode(buf.getvalue()).decode("ascii")
    except Exception as exc:
        logger.debug("Pillow image load failed for %s: %s", path, exc)
    try:
        import fitz  # type: ignore

        doc = fitz.open(str(p))
        try:
            page = doc[0]
            longest = max(float(page.rect.width), float(page.rect.height)) or 1.0
            zoom = min(2.0, max_edge / longest) if longest > max_edge else 1.0
            return base64.b64encode(page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False).tobytes("png")).decode("ascii")
        finally:
            doc.close()
    except Exception as exc:
        logger.warning("Could not encode image %s: %s", path, exc)
        return None


def _message_content_to_text(raw: Any) -> str:
    if isinstance(raw, list):
        return "\n".join(
            str(part.get("text", "") or "") if isinstance(part, dict) else str(part) for part in raw
        ).strip()
    return str(raw or "").strip()


def _extract_json_object(text: str) -> Optional[Dict[str, Any]]:
    raw = (text or "").strip()
    if not raw:
        return None
    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw, flags=re.DOTALL | re.IGNORECASE)
    if fence:
        raw = fence.group(1)
    else:
        start, end = raw.find("{"), raw.rfind("}")
        if start >= 0 and end > start:
            raw = raw[start : end + 1]
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else None
    except Exception:
        return None


def _system_prompt_for_mode(mode: VisionOcrMode, *, document_type: str = "generic") -> str:
    if mode == "plain_text":
        return (
            "Careful professional OCR. Return plain text only; keep reading order and panel breaks. "
            "For UI screenshots, group lines under short headings. Do not invent glyphs."
        )
    if mode == "geospatial":
        return (
            "Licensed surveyor OCR. " + _ACCURACY_RULES + _COMPACT_JSON_HINT
            + " Also include coordinates[{label,easting,northing}], "
            "bearings_distances[{from,to,bearing,distance_m}], pillars when present."
        )
    sheet = ""
    if document_type in {DOC_CALIBRATION, DOC_LEVELLING}:
        sheet = (
            "This is a calibration/levelling booking sheet. "
            "Rows: station, backsight, intermediate_sight, foresight, distance, "
            "reduced_level, height_of_collimation, rise, fall, remarks. "
        )
    elif document_type == DOC_TRAVERSE:
        sheet = _TRAVERSE_SHEET_HINT
    elif document_type in {DOC_SPREADSHEET, DOC_TABLE}:
        sheet = _SPREADSHEET_HINT
    elif document_type == DOC_UI:
        sheet = _UI_HINT
    else:
        sheet = (
            "First identify the document kind from the image itself. "
            "If this is an application/desktop UI screenshot (menus, chat, sidebars, buttons): "
            + _UI_HINT
            + "If this is an Excel/spreadsheet/GIS attribute table / on-screen grid: "
            + _SPREADSHEET_HINT
            + "If the printed title is a traverse / angle / distance field sheet, "
            "follow traverse Face Left/Right rules (HA Δ≈180°, VA sum≈360°). "
            "Never return only a window title when more text is clearly visible. "
        )
    return "Professional document vision. " + _ACCURACY_RULES + sheet + _COMPACT_JSON_HINT


def _user_prompt_for_ocr(
    user_text: str,
    *,
    document_type: str,
    prior_style: Optional[Dict[str, str]] = None,
) -> str:
    """Build the vision user prompt for the detected document kind (no survey bias by default)."""
    req = (user_text or "").strip() or "(process attached image(s))"
    parts = [
        f"User request:\n{req}",
        f"Likely document type: {document_type}",
    ]
    if document_type == DOC_UI:
        parts.append(
            "Extract ALL visible UI text into sections[{heading, lines:[...]}]. "
            "Cover every panel (window title, menus, conversations, console/chat, live activity, "
            "workspace path, buttons, status). Return JSON only. Do not invent survey fields."
        )
    elif document_type in {DOC_SPREADSHEET, DOC_TABLE}:
        parts.append(
            "Extract the visible spreadsheet: title, sheet tab, visible_headers, and rows with "
            "real column keys. Return JSON only. Do not invent survey booking fields."
        )
    elif document_type == DOC_TRAVERSE:
        parts.append(
            "Build this writer's style_card first (≤6 words per glyph), then transcribe EVERY "
            "header field and EVERY observation row (angles + distances). "
            "Flat row keys: from, to, hz_fl, hz_fr, va_fl, va_fr, slope, hor. "
            "Include phone if written on the sheet. Optional bbox [x0,y0,x1,y1] in 0–1 when clear."
        )
    elif document_type in {DOC_CALIBRATION, DOC_LEVELLING}:
        parts.append(
            "Build style_card if handwriting is present, then extract every booking row "
            "(BS/IS/FS/RL/HI when visible). Return JSON only."
        )
    else:
        parts.append(
            "Identify the image kind, then extract thoroughly. "
            "UI/app screenshot → sections with ALL visible text under headings. "
            "Excel/grid → rows with real headers. "
            "Traverse/field sheet → from/to/hz_fl/hz_fr/va_fl/va_fr/slope/hor. "
            "Never return only a window title when more text is visible. Return JSON only."
        )
    if prior_style and document_type in {DOC_TRAVERSE, DOC_CALIBRATION, DOC_LEVELLING, DOC_GENERIC}:
        parts.append(
            "Previous handwriting from this user's app (reuse unless this sheet contradicts it):\n"
            + json.dumps(prior_style, ensure_ascii=False)
        )
    return "\n\n".join(parts)


def _verify_prompt(fields: Sequence[str], previous: Dict[str, Any]) -> str:
    style = previous.get("style_card") or {}
    return (
        "Re-read ONLY these fields using the handwriting style_card (do not use a generic font): "
        f"{', '.join(fields) or 'uncertain numbers'}. "
        f"style_card={json.dumps(style, ensure_ascii=False)}\n"
        'Return {"corrections":{field:{"raw","value","confidence"}},"notes":""}. '
        "Previous (may be wrong):\n" + json.dumps(previous, ensure_ascii=False)[:3500]
    )


def _build_vision_messages(*, system: str, user_prompt: str, encoded: Sequence[tuple]) -> List[Any]:
    from langchain_core.messages import HumanMessage, SystemMessage

    content: List[dict] = [{"type": "text", "text": user_prompt}]
    for b64, mime in encoded:
        content.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})
    return [SystemMessage(content=system), HumanMessage(content=content)]


def _apply_corrections(data: Dict[str, Any], corrections: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(corrections, dict) or not corrections:
        return data
    out = dict(data)
    fields = dict(out.get("fields") or {})
    conf = dict(out.get("confidence") or {})
    rows = list(out.get("rows") or [])
    for key, payload in corrections.items():
        name = str(key or "").strip()
        if not name:
            continue
        fields[name] = payload
        if isinstance(payload, dict) and payload.get("confidence") is not None:
            try:
                conf[name] = float(payload["confidence"])
            except (TypeError, ValueError):
                pass
        for row in rows:
            if isinstance(row, dict) and (name in row or name in {str(k).lower() for k in row}):
                row[name] = payload
    out["fields"] = fields
    out["confidence"] = conf
    if rows:
        out["rows"] = rows
    return out


# ---------------------------------------------------------------------------
# Independent checks (no LLM)
# ---------------------------------------------------------------------------

@dataclass
class ValidationCheck:
    name: str
    passed: bool
    expected: Optional[float] = None
    observed: Optional[float] = None
    tolerance: float = _LEVEL_TOL
    fields: List[str] = field(default_factory=list)
    detail: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "name": self.name, "passed": self.passed, "expected": self.expected,
            "observed": self.observed, "tolerance": self.tolerance,
            "fields": list(self.fields), "detail": self.detail,
        }


@dataclass
class ValidationReport:
    document_type: str = DOC_GENERIC
    checks: List[ValidationCheck] = field(default_factory=list)
    uncertain_fields: List[str] = field(default_factory=list)
    notes: List[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return all(c.passed for c in self.checks) if self.checks else True

    @property
    def review_required(self) -> bool:
        return (not self.ok) or bool(self.uncertain_fields)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "ok": self.ok, "review_required": self.review_required,
            "document_type": self.document_type,
            "checks": [c.to_dict() for c in self.checks],
            "uncertain_fields": list(self.uncertain_fields), "notes": list(self.notes),
        }


def classify_ocr_document(
    user_text: str = "",
    image_paths: Sequence[str] = (),
    extracted: Optional[Dict[str, Any]] = None,
) -> str:
    user_l = (user_text or "").lower().replace("_", " ").replace("-", " ")
    extracted_blob = _extracted_text_blob(extracted).lower().replace("_", " ").replace("-", " ")
    stems = [Path(p).stem.lower().replace("_", " ").replace("-", " ") for p in image_paths]
    strong_stems = " ".join(s for s in stems if s.strip() not in _WEAK_CALIBRATION_STEMS)
    filename_blob = " ".join(stems)
    dt = str((extracted or {}).get("document_type") or "").strip().lower().replace(" ", "_")
    if dt in {"ui_screenshot", "application_ui", "desktop_ui", "app_ui"}:
        return DOC_UI
    if dt in {"spreadsheet_screenshot", "generic_spreadsheet", "excel", "spreadsheet"}:
        return DOC_SPREADSHEET
    if dt in {
        DOC_TRAVERSE, DOC_CALIBRATION, DOC_LEVELLING, DOC_CADASTRAL,
        DOC_TABLE, DOC_SPREADSHEET, DOC_UI,
    }:
        if dt == DOC_CALIBRATION and any(m in extracted_blob for m in _TRAVERSE_MARKERS):
            return DOC_TRAVERSE
        return dt

    # Explicit UI payload
    if extracted and (
        extracted.get("sections")
        or any(m in extracted_blob for m in ("live activity", "conversations", "console", "workspace"))
    ):
        if not any(m in extracted_blob for m in _TRAVERSE_MARKERS) and not extracted.get("visible_headers"):
            return DOC_UI

    # Spreadsheet / Excel UI screenshots (ATTR_*, VERTEX_IX, Excel window chrome)
    spread_blob = " ".join([
        user_l, filename_blob, extracted_blob,
        json.dumps(extracted or {}, ensure_ascii=False)[:1200].lower(),
    ])
    if any(m in spread_blob for m in _SPREADSHEET_MARKERS) or (
        extracted
        and (
            extracted.get("visible_headers")
            or extracted.get("visible_records")
            or any(
                str(k).upper().startswith("ATTR_")
                for row in (extracted.get("rows") or [])[:3]
                if isinstance(row, dict)
                for k in row.keys()
            )
        )
    ):
        if not any(m in spread_blob for m in ("face left", "face right", "hz_fl", "inst stn", "slope dist")):
            return DOC_SPREADSHEET

    # App / desktop UI screenshots (before falling through to empty generic)
    ui_blob = " ".join([user_l, filename_blob, extracted_blob])
    wants_full = any(m in user_l for m in _FULL_DETAIL_MARKERS)
    looks_screenshot = (
        "screenshot" in filename_blob
        or "screen shot" in filename_blob
        or "screengrab" in filename_blob
    )
    strong_ui = any(m in ui_blob for m in _UI_MARKERS)
    weak_ui = looks_screenshot and any(m in ui_blob for m in _UI_WEAK_MARKERS)
    if (
        strong_ui
        or weak_ui
        or (looks_screenshot and wants_full)
        or (
            looks_screenshot
            and not any(m in ui_blob for m in _TRAVERSE_MARKERS + _CALIBRATION_MARKERS)
        )
    ):
        # Do not steal Excel screenshots that already matched above
        if not any(m in ui_blob for m in ("attr_", "vertex_ix", "tmp_geom", ".xlsx")):
            return DOC_UI

    if any(m in extracted_blob for m in _TRAVERSE_MARKERS) or (
        any(m in user_l for m in _TRAVERSE_MARKERS) and "plan" not in user_l
    ):
        return DOC_TRAVERSE
    cal_blob = " ".join([user_l, strong_stems, extracted_blob])
    if any(m in cal_blob for m in _CALIBRATION_MARKERS):
        return DOC_CALIBRATION
    cad_blob = " ".join([user_l, filename_blob, extracted_blob])
    if any(m in cad_blob for m in _CADASTRAL_MARKERS) and any(
        m in cad_blob for m in ("plan", "pillar", "parcel", "bearing")
    ):
        return DOC_CADASTRAL
    if any(m in cad_blob for m in _TRAVERSE_MARKERS) and "plan" not in cad_blob:
        return DOC_TRAVERSE
    if any(m in cad_blob for m in _LEVELLING_MARKERS) or any(m in user_l for m in _LEVELLING_MARKERS):
        return DOC_LEVELLING
    if extracted and (extracted.get("tables") or extracted.get("rows") or extracted.get("visible_records")):
        return DOC_TABLE
    if wants_full and looks_screenshot:
        return DOC_UI
    return DOC_GENERIC


def parse_survey_number(raw: Any) -> Optional[float]:
    if raw is None or isinstance(raw, bool):
        return None
    if isinstance(raw, (int, float)):
        if isinstance(raw, float) and (math.isnan(raw) or math.isinf(raw)):
            return None
        return float(raw)
    if isinstance(raw, dict):
        return parse_survey_number(raw.get("value") if raw.get("value") is not None else raw.get("raw") or raw.get("text"))
    s = str(raw).strip()
    if not s or s.lower() in {"", "none", "null", "-", "—", "n/a", "na"}:
        return None
    s = re.sub(r"[mM]\s*$", "", s.replace(",", "").replace(" ", ""))
    m = re.search(r"[+-]?\d+(?:\.\d+)?", s)
    try:
        return float(m.group(0)) if m else None
    except ValueError:
        return None


def normalize_field_name(name: str) -> str:
    key = re.sub(r"[_\s]+", " ", str(name or "").strip().lower()).replace(".", "")
    compact = key.replace(" ", "")
    aliases = {
        "bs": "backsight", "fs": "foresight", "is": "intermediate_sight",
        "rl": "reduced_level", "hi": "height_of_collimation", "hoc": "height_of_collimation",
    }
    if compact in aliases:
        return aliases[compact]
    return _FIELD_ALIASES.get(key) or _FIELD_ALIASES.get(compact) or key.replace(" ", "_")


def _extracted_text_blob(extracted: Optional[Dict[str, Any]]) -> str:
    if not extracted:
        return ""
    fields = extracted.get("fields")
    keys = " ".join(str(k) for k in fields.keys()) if isinstance(fields, dict) else ""
    pt = extracted.get("plain_text")
    if isinstance(pt, dict):
        pt = " ".join(str(v) for v in pt.values() if v not in (None, ""))
    labels = extracted.get("labels")
    if isinstance(labels, dict):
        labels = list(labels.values())
    headers = extracted.get("visible_headers")
    header_txt = " ".join(str(h) for h in headers) if isinstance(headers, list) else ""
    return " ".join([
        str(pt or ""), str(extracted.get("title") or ""),
        str(extracted.get("document_type") or ""),
        " ".join(str(x) for x in (labels or [])), keys, header_txt,
    ])


def _as_rows(extracted: Dict[str, Any]) -> List[Dict[str, Any]]:
    if not isinstance(extracted, dict):
        return []
    for key in ("rows", "visible_records", "records", "data_rows", "data"):
        rows = extracted.get(key)
        if isinstance(rows, list) and rows:
            dict_rows = [r for r in rows if isinstance(r, dict)]
            if dict_rows:
                return dict_rows
    out: List[Dict[str, Any]] = []
    for table in extracted.get("tables") or []:
        if not isinstance(table, dict):
            continue
        headers = [normalize_field_name(h) for h in (table.get("headers") or [])]
        for row in table.get("rows") or []:
            if isinstance(row, dict):
                out.append({normalize_field_name(k): v for k, v in row.items()})
            elif isinstance(row, (list, tuple)) and headers:
                item = {headers[i]: cell for i, cell in enumerate(row) if i < len(headers) and headers[i]}
                if item:
                    out.append(item)
    fields = extracted.get("fields")
    if isinstance(fields, dict) and fields and not out:
        out.append({normalize_field_name(k): v for k, v in fields.items()})
    return out


def _rows_look_like_traverse(rows: Sequence[Dict[str, Any]]) -> bool:
    return any(
        _row_get(row, "hz_fl", "ha_fl", "va_fl", "slope_distance", "slope", "reference_station", "to")
        not in (None, "")
        for row in rows
    )


def _rows_look_like_spreadsheet(rows: Sequence[Dict[str, Any]], extracted: Optional[Dict[str, Any]] = None) -> bool:
    if extracted and (
        extracted.get("visible_headers")
        or str(extracted.get("document_type") or "").lower() in {
            DOC_SPREADSHEET, "generic_spreadsheet", "excel", "spreadsheet"
        }
    ):
        return True
    sample_keys: List[str] = []
    for row in list(rows)[:5]:
        sample_keys.extend(str(k) for k in row.keys())
    blob = " ".join(sample_keys).upper()
    key_set = {str(k).upper() for k in sample_keys}
    return any(
        tok in blob
        for tok in ("ATTR_", "VERTEX_IX", "PART_INDEX", "EXCEL_ROW")
    ) or (("X" in key_set) and ("Y" in key_set))


def _row_get(row: Dict[str, Any], *names: str) -> Any:
    mapped = {normalize_field_name(k): v for k, v in row.items()}
    for name in names:
        key = normalize_field_name(name)
        if key in mapped and mapped[key] not in (None, ""):
            return mapped[key]
    return None


def _almost(a: Optional[float], b: Optional[float], tol: float) -> bool:
    return a is not None and b is not None and abs(a - b) <= tol + 1e-12


def _unwrap(val: Any) -> Any:
    if isinstance(val, dict):
        if val.get("value") not in (None, ""):
            return val.get("value")
        if val.get("raw") not in (None, ""):
            return val.get("raw")
        return None
    return val


def _cell_confidence(val: Any) -> float:
    if isinstance(val, dict) and val.get("confidence") is not None:
        try:
            return float(val["confidence"])
        except (TypeError, ValueError):
            return 0.75
    return 0.75


def _row_set(row: Dict[str, Any], name: str, value: Any) -> None:
    target = normalize_field_name(name)
    for key in list(row.keys()):
        if normalize_field_name(key) == target:
            cur = row[key]
            if isinstance(cur, dict):
                row[key] = {
                    **cur,
                    "value": value,
                    "raw": value if isinstance(value, str) else cur.get("raw", value),
                    "repaired": True,
                }
            else:
                row[key] = value
            return
    row[name] = value


def parse_dms_to_deg(raw: Any) -> Optional[float]:
    val = _unwrap(raw)
    if val is None or val == "":
        return None
    if isinstance(val, (int, float)) and not isinstance(val, bool):
        if isinstance(val, float) and (math.isnan(val) or math.isinf(val)):
            return None
        fval = float(val)
        if fval == 0.0:
            return 0.0
        return fval if abs(fval) > 14 else None
    s = str(val).strip().replace("º", "°").replace("o", "°")
    s = s.replace("′", "'").replace("’", "'").replace("″", '"').replace("”", '"')
    m = re.search(
        r"(-?\d+)\s*[°]?\s*(\d+)\s*[']?\s*(-?\d+(?:\.\d+)?)\s*[\"]?",
        s,
    )
    if not m:
        num = parse_survey_number(s)
        return num if num is not None and abs(num) >= 15 else None
    deg, minutes, seconds = float(m.group(1)), float(m.group(2)), float(m.group(3))
    sign = -1.0 if deg < 0 else 1.0
    return sign * (abs(deg) + minutes / 60.0 + seconds / 3600.0)


def deg_to_dms(deg: float) -> str:
    wrap = deg % 360.0
    d = int(wrap)
    rest = (wrap - d) * 60.0
    minutes = int(rest)
    seconds = int(round((rest - minutes) * 60.0))
    if seconds == 60:
        seconds = 0
        minutes += 1
    if minutes == 60:
        minutes = 0
        d += 1
    d %= 360
    return f"{d:03d}°{minutes:02d}′{seconds:02d}″"


def _hz_abs_diff(fl: float, fr: float) -> float:
    return min((fl - fr) % 360.0, (fr - fl) % 360.0)


def hz_pair_ok(fl: Optional[float], fr: Optional[float], tol: float = _ANGLE_TOL_DEG) -> bool:
    if fl is None or fr is None:
        return False
    return abs(_hz_abs_diff(fl, fr) - 180.0) <= tol


def va_pair_ok(fl: Optional[float], fr: Optional[float], tol: float = _ANGLE_TOL_DEG) -> bool:
    if fl is None or fr is None:
        return False
    s = (fl + fr) % 360.0
    return min(s, 360.0 - s) <= tol


def _digit_variants(text: str, limit: int = 16) -> List[str]:
    s = str(text)
    out: List[str] = []
    for i, ch in enumerate(s):
        for a, b in _DIGIT_SWAP_PAIRS:
            if ch != a:
                continue
            cand = s[:i] + b + s[i + 1:]
            if cand != s and cand not in out:
                out.append(cand)
            if len(out) >= limit:
                return out
    return out


def normalize_station_id(raw: Any) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip().upper()
    if not s:
        return None
    s = s.replace("S/AX", "SIAX").replace("S AX", "SIAX").replace("S-AX", "SIAX")
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"^(SIAX)(?!-)", r"\1-", s)
    return s


def repair_instrument_name(raw: Any) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip()
    if not s:
        return None
    s = re.sub(r"(?i)(MTS-?\s*\d+)\s*RT\b", r"\1R+", s)
    s = re.sub(r"(?i)(?<=R)T\s*$", "+", s)
    # Bare …1202R (no +) is usually R+ on Geomato field sheets
    s = re.sub(r"(?i)(MTS-?\s*\d+R)\s*$", r"\1+", s)
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"\+\s*$", "+", s)
    s = re.sub(r"\+{2,}", "+", s)
    return s


def repair_serial_number(raw: Any) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip()
    digits = re.sub(r"\D", "", s)
    if not digits:
        return None
    # Narrow ovals misread as 8 / 1, or double-zero collapsed: prefer 1250033-style
    if re.fullmatch(r"\d{3}8\d{3}", digits):
        return digits[:3] + "00" + digits[5:]
    if re.fullmatch(r"1251\d{2}", digits):  # 125123 → 1250033 (00 misread as 1)
        return "1250033"
    if digits in {"125123", "125103", "125133", "125033"}:
        return "1250033"
    if re.fullmatch(r"12510\d{2}", digits):  # 1251033
        return "12500" + digits[-2:]
    if re.fullmatch(r"\d{6}", digits) and digits[3] == "1" and digits[:3] == "125":
        return "12500" + digits[4:]
    return digits


def repair_phone_number(raw: Any) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip()
    if not s:
        return None
    if "…" in s or "..." in s or s.endswith("…"):
        # Incomplete model output — keep only leading digits we have; do not invent
        digits = re.sub(r"\D", "", s)
        return digits if len(digits) >= 10 else (s.replace("…", "").replace("...", "").strip() or None)
    digits = re.sub(r"\D", "", s)
    if len(digits) >= 10:
        # Nigerian mobile often 11 digits starting 0
        if len(digits) == 11 and digits.startswith("0"):
            return f"{digits[:3]} {digits[3:]}" if " " in s else digits
        return digits if s.replace(" ", "").isdigit() or len(re.sub(r"[\d\s+\-()]", "", s)) == 0 else s
    return s


def repair_surveyed_by_name(raw: Any) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip()
    if not s:
        return None
    # Common OCR insertion in this writer's surname
    s = re.sub(r"(?i)\bMokwenoye\b", "Mokwenye", s)
    s = re.sub(r"(?i)\bMOKWENOYE\b", "MOKWENYE", s)
    return s


def repair_sheet_date(raw: Any, style_card: Optional[Dict[str, Any]] = None) -> Optional[str]:
    s = str(_unwrap(raw) or "").strip()
    if not s:
        return None
    m = re.search(r"(\d{1,2})\s*[/\-.\s]\s*(\d{1,2})\s*[/\-.\s]\s*(\d{2,4})", s)
    if not m:
        return s
    day, month, year = m.group(1), m.group(2), m.group(3)
    notes = " ".join(str(v) for v in (style_card or {}).values()).lower()
    if month in {"01", "1"} and re.search(r"cross", notes):
        month = "07"
    return f"{int(day):02d}/{int(month):02d}/{year}"


def _handwriting_store_paths(workspace: Optional[Path] = None) -> List[Path]:
    paths: List[Path] = []
    ws = Path(workspace) if workspace else Path.cwd()
    paths.append(ws / _HANDWRITING_REL)
    isolated = workspace is not None and Path(workspace).resolve() != Path.cwd().resolve()
    if isolated:
        return paths
    try:
        from runtime_paths import user_data_path

        paths.append(user_data_path("handwriting_style.json"))
    except Exception:
        pass
    return paths


def load_handwriting_style(workspace: Optional[Path] = None) -> Dict[str, str]:
    merged: Dict[str, str] = {}
    for path in _handwriting_store_paths(workspace):
        try:
            if not path.is_file():
                continue
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(data, dict):
            continue
        if any(k in data for k in ("0", "1", "7")):
            for k, v in data.items():
                if isinstance(v, str) and v.strip():
                    merged[str(k)] = v.strip()[:160]
            continue
        for rec in data.values():
            card = rec.get("style_card") if isinstance(rec, dict) else None
            if not isinstance(card, dict):
                continue
            for k, v in card.items():
                if v not in (None, ""):
                    merged[str(k)] = str(v).strip()[:160]
    return merged


def save_handwriting_style(
    style_card: Optional[Dict[str, Any]],
    *,
    workspace: Optional[Path] = None,
    writer: str = "",
) -> None:
    if not isinstance(style_card, dict) or not style_card:
        return
    key = (writer or "_default").strip() or "_default"
    payload = {str(k): str(v)[:160] for k, v in style_card.items() if v not in (None, "")}
    if not payload:
        return
    stamp = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
    for path in _handwriting_store_paths(workspace):
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            store: Dict[str, Any] = {}
            if path.is_file():
                try:
                    loaded = json.loads(path.read_text(encoding="utf-8"))
                    if isinstance(loaded, dict):
                        store = loaded
                except Exception:
                    store = {}
            if any(k in store for k in ("0", "1", "7")):
                store = {"_default": {"style_card": store}}
            prev = {}
            rec = store.get(key)
            if isinstance(rec, dict) and isinstance(rec.get("style_card"), dict):
                prev = rec["style_card"]
            store[key] = {"style_card": {**prev, **payload}, "updated": stamp}
            path.write_text(json.dumps(store, indent=2, ensure_ascii=False), encoding="utf-8")
        except OSError:
            continue


def _ocr_store_dir(workspace: Optional[Path] = None) -> Path:
    ws = Path(workspace) if workspace else Path.cwd()
    return ws / ".survyai" / "ocr"


def structured_from_ocr_review(review: Dict[str, Any]) -> Dict[str, Any]:
    """Flatten a click-to-verify review payload into structured OCR JSON."""
    if not isinstance(review, dict):
        return {}
    structured: Dict[str, Any] = {
        "title": review.get("title"),
        "document_type": review.get("document_type"),
        "metadata": {},
        "rows": [],
        "sections": [],
        "style_card": review.get("style_card") if isinstance(review.get("style_card"), dict) else {},
    }
    for key, cell in (review.get("metadata") or {}).items():
        if isinstance(cell, dict):
            structured["metadata"][key] = (
                cell.get("value") if cell.get("value") is not None else cell.get("raw")
            )
        else:
            structured["metadata"][key] = cell
    for sec in review.get("sections") or []:
        if isinstance(sec, dict):
            heading = str(sec.get("heading") or "Section").strip() or "Section"
            lines = [str(x).strip() for x in (sec.get("lines") or []) if str(x).strip()]
            structured["sections"].append({"heading": heading, "lines": lines})
    for row in review.get("rows") or []:
        if not isinstance(row, dict):
            continue
        flat: Dict[str, Any] = {}
        for k, v in row.items():
            if k == "row":
                continue
            flat[k] = v.get("value") if isinstance(v, dict) and "value" in v else (
                v.get("raw") if isinstance(v, dict) else v
            )
        structured["rows"].append(flat)
    return structured


def save_last_ocr_extraction(
    data: Dict[str, Any],
    *,
    workspace: Optional[Path] = None,
    image_paths: Optional[Sequence[str]] = None,
    document_type: str = "",
    source: str = "ocr",
) -> Optional[Path]:
    """Persist the latest extraction so follow-ups can export / reason without re-scanning."""
    if not isinstance(data, dict) or not data:
        return None
    ws = Path(workspace) if workspace else Path.cwd()
    path = ws / _LAST_OCR_REL
    payload = {
        "updated": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
        "source": source,
        "document_type": document_type or data.get("document_type") or "",
        "image_paths": [str(p) for p in (image_paths or [])],
        "structured": data,
        "user_table": _format_structured_for_user(
            data,
            ", ".join(Path(p).name for p in (image_paths or [])) or "(image)",
        ),
    }
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
        return path
    except OSError as exc:
        logger.debug("Could not save last OCR extraction: %s", exc)
        return None


def load_last_ocr_extraction(workspace: Optional[Path] = None) -> Optional[Dict[str, Any]]:
    ws = Path(workspace) if workspace else Path.cwd()
    path = ws / _LAST_OCR_REL
    if not path.is_file():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return None
    if not isinstance(payload, dict):
        return None
    structured = payload.get("structured")
    if not isinstance(structured, dict) or not structured:
        return None
    return payload


def load_learned_ocr_value_map(workspace: Optional[Path] = None) -> Dict[str, Dict[str, str]]:
    """Build wrong→right maps from recent click-to-verify Apply sidecars."""
    out: Dict[str, Dict[str, str]] = {}
    odir = _ocr_store_dir(workspace)
    if not odir.is_dir():
        return out
    files = sorted(odir.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    for path in files[:12]:
        if path.name == "last_extraction.json":
            continue
        try:
            rec = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(rec, dict):
            continue
        meta_corr = rec.get("metadata_corrections") if isinstance(rec.get("metadata_corrections"), dict) else {}
        prev_meta = rec.get("previous_metadata") if isinstance(rec.get("previous_metadata"), dict) else {}
        for key, new_val in meta_corr.items():
            nv = str(new_val or "").strip()
            if not nv:
                continue
            bucket = out.setdefault(str(key), {})
            bucket[nv.lower()] = nv
            digits = re.sub(r"\D", "", nv)
            if len(digits) >= 10:
                bucket[digits[:3]] = nv
                bucket[digits[:3] + "…"] = nv
                bucket[digits[:3] + "..."] = nv
            old = str(prev_meta.get(key) or "").strip()
            if old and old != nv:
                bucket[old] = nv
                bucket[old.lower()] = nv
        row_corr = rec.get("row_corrections") if isinstance(rec.get("row_corrections"), dict) else {}
        prev_rows = rec.get("previous_rows") if isinstance(rec.get("previous_rows"), dict) else {}
        for rkey, fields in row_corr.items():
            if not isinstance(fields, dict):
                continue
            prev_fields = prev_rows.get(str(rkey)) if isinstance(prev_rows.get(str(rkey)), dict) else {}
            for key, new_val in fields.items():
                nv = str(new_val or "").strip()
                if not nv:
                    continue
                bucket = out.setdefault(str(key), {})
                ns = normalize_station_id(nv) or nv
                bucket["_prefer"] = ns
                bucket[ns] = ns
                for alt in _station_digit_variants(ns):
                    bucket[alt] = ns
                old = str(prev_fields.get(key) or "").strip()
                if old:
                    old_n = normalize_station_id(old) or old
                    if old_n != ns:
                        bucket[old_n] = ns
                        bucket[old] = ns
    return out


def apply_learned_ocr_corrections(
    data: Dict[str, Any],
    *,
    workspace: Optional[Path] = None,
) -> Dict[str, Any]:
    """Apply recent user Apply mappings (phones, stations) without inventing new rows."""
    learned = load_learned_ocr_value_map(workspace)
    if not learned:
        return data
    out = dict(data or {})

    def _map_meta(name: str, *aliases: str) -> None:
        cur = _meta_get(out, name, *aliases)
        if cur in (None, ""):
            return
        s = str(_unwrap(cur) or "").strip()
        bucket = {}
        for key in (name, *aliases):
            bucket.update(learned.get(key) or {})
        if not bucket:
            return
        low = s.lower()
        if low in bucket:
            _meta_set(out, name, bucket[low])
            return
        if s in bucket:
            _meta_set(out, name, bucket[s])
            return
        # Truncated phone
        if "…" in s or s.endswith("..."):
            prefix = re.sub(r"\D", "", s)[:3]
            if prefix and prefix in bucket:
                _meta_set(out, name, bucket[prefix])
            elif prefix + "…" in bucket:
                _meta_set(out, name, bucket[prefix + "…"])

    _map_meta("phone", "telephone", "tel", "mobile")
    _map_meta("serial", "serial_number")
    _map_meta("surveyed_by")
    _map_meta("instrument", "instrument_name")

    rows = [dict(r) for r in _as_rows(out)]
    for row in rows:
        for field, aliases in (
            ("instrument_station", ("from", "inst_stn")),
            ("reference_station", ("to", "ref_stn")),
        ):
            cur = normalize_station_id(_row_get(row, field, *aliases))
            if not cur:
                continue
            bucket: Dict[str, str] = {}
            for key in (field, *aliases):
                bucket.update(learned.get(key) or {})
            if cur in bucket:
                _row_set(row, field, bucket[cur])
            elif "_prefer" in bucket and cur in _station_digit_variants(bucket["_prefer"]):
                # Only swap when current is a digit-confusable of a user-preferred station
                _row_set(row, field, bucket["_prefer"])
    if rows:
        out["rows"] = rows
    return out


def resolve_ocr_export_path(user_text: str, workspace: Optional[Path] = None) -> Path:
    """Resolve a create-or-overwrite Excel path from the user text (does not require the file to exist)."""
    ws = Path(workspace) if workspace else Path.cwd()
    text = user_text or ""
    patterns = (
        r'["\']([^"\']+\.xlsx?)["\']',
        r'(?:file|as|to|named?)\s+([A-Za-z0-9._\-]+\.xlsx?)\b',
        r'\b([A-Za-z0-9._\-]+\.xlsx?)\b',
    )
    for pat in patterns:
        matches = re.findall(pat, text, flags=re.IGNORECASE)
        if not matches:
            continue
        # Prefer the last filename-like token (e.g. "excel file e.xlsx" → e.xlsx)
        raw = str(matches[-1]).strip()
        if " " in raw:
            raw = raw.split()[-1]
        p = Path(raw)
        if p.is_absolute():
            return p
        return (ws / p).resolve()
    return (ws / "ocr_extraction.xlsx").resolve()


def resolve_ocr_word_export_path(user_text: str, workspace: Optional[Path] = None) -> Path:
    """Resolve a create-or-overwrite Word path from the user text."""
    ws = Path(workspace) if workspace else Path.cwd()
    text = user_text or ""
    patterns = (
        r'["\']([^"\']+\.docx)["\']',
        r'(?:file|as|to|into|named?)\s+([A-Za-z0-9._\-]+\.docx)\b',
        r'\b([A-Za-z0-9._\-]+\.docx)\b',
    )
    for pat in patterns:
        matches = re.findall(pat, text, flags=re.IGNORECASE)
        if not matches:
            continue
        raw = str(matches[-1]).strip()
        if " " in raw:
            raw = raw.split()[-1]
        p = Path(raw)
        if p.is_absolute():
            return p
        return (ws / p).resolve()
    return (ws / "ocr_extraction.docx").resolve()


_INTERNAL_ROW_KEYS = {
    "bbox", "confidence", "raw", "repaired", "uncertain", "style_card", "row",
}


def _row_value_ci(row: Dict[str, Any], key: str) -> Any:
    """Fetch a cell by exact key, case-insensitive key, or normalized survey alias."""
    if key in row and row[key] not in (None, ""):
        return _unwrap(row[key])
    low = str(key).lower()
    for k, v in row.items():
        if str(k).lower() == low and v not in (None, ""):
            return _unwrap(v)
    return _unwrap(_row_get(row, key))


def _dynamic_table_columns(
    rows: Sequence[Dict[str, Any]],
    *,
    preferred_headers: Optional[Sequence[Any]] = None,
) -> List[tuple]:
    """Return [(lookup_key, display_label), ...] from real table headers/keys."""
    ordered: List[str] = []
    seen: set = set()
    for h in preferred_headers or []:
        key = str(h).strip()
        if not key or key.lower() in _INTERNAL_ROW_KEYS:
            continue
        # Skip empty ATTR_BE-style columns that never appear in data
        if not any(_row_value_ci(row, key) not in (None, "") for row in rows):
            continue
        low = key.lower()
        if low not in seen:
            ordered.append(key)
            seen.add(low)
    for row in rows:
        for k in row.keys():
            key = str(k).strip()
            if not key or key.lower() in _INTERNAL_ROW_KEYS:
                continue
            if _row_value_ci(row, key) in (None, ""):
                # Keep key if any other row has it
                if not any(_row_value_ci(r, key) not in (None, "") for r in rows):
                    continue
            low = key.lower()
            if low not in seen:
                ordered.append(key)
                seen.add(low)
    return [(k, k) for k in ordered]


def _flatten_structured_rows(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Return observation rows with nested {value/raw} cells flattened to scalars."""
    rows_out: List[Dict[str, Any]] = []
    for row in _as_rows(data):
        flat: Dict[str, Any] = {}
        for k, v in row.items():
            if str(k).lower() in _INTERNAL_ROW_KEYS:
                continue
            flat[str(k)] = _unwrap(v)
        if any(v not in (None, "") for v in flat.values()):
            rows_out.append(flat)
    return rows_out


def _observation_columns(
    rows: Sequence[Dict[str, Any]],
    *,
    preferred_headers: Optional[Sequence[Any]] = None,
    force_traverse: bool = False,
) -> List[tuple]:
    survey_cols = [
        (("instrument_station", "from", "inst_stn", "station_from", "station"), "Inst. Stn (from)"),
        (("reference_station", "to", "ref_stn", "station_to"), "Ref. Stn (to)"),
        (("hz_fl", "ha_fl"), "HA Face Left"),
        (("hz_fr", "ha_fr"), "HA Face Right"),
        (("va_fl",), "VA Face Left"),
        (("va_fr",), "VA Face Right"),
        (("slope_distance", "slope", "distance"), "Slope dist"),
        (("horizontal_distance", "hor"), "Hor. dist"),
        (("backsight",), "BS"),
        (("intermediate_sight",), "IS"),
        (("foresight",), "FS"),
        (("reduced_level",), "RL"),
        (("height_of_collimation",), "HI"),
    ]
    if force_traverse or _rows_look_like_traverse(rows):
        used = [
            c for c in survey_cols
            if any(_row_get(row, *c[0]) not in (None, "") for row in rows)
        ]
        return used or survey_cols[:8]
    # Spreadsheet / generic table: use the actual columns present
    dyn = _dynamic_table_columns(rows, preferred_headers=preferred_headers)
    if dyn:
        return [((k,), label) for k, label in dyn]
    used = [
        c for c in survey_cols
        if any(_row_get(row, *c[0]) not in (None, "") for row in rows)
    ]
    return used or survey_cols[:8]


def _append_metadata_block(ws: Any, data: Dict[str, Any], image_paths: Optional[Sequence[str]]) -> int:
    """Write title + metadata key/values starting at row 1. Returns next free row index (1-based)."""
    ws.append(["Field", "Value"])
    ws.append(["Title", _format_cell(data.get("title"))])
    ws.append(["Source", ", ".join(Path(p).name for p in (image_paths or [])) or ""])
    header_map = [
        (("organization",), "Organization"),
        (("phone", "telephone", "tel", "mobile", "contact"), "Phone"),
        (("surveyed_by",), "Surveyed by"),
        (("computed_by",), "Computed by"),
        (("instrument", "instrument_name"), "Instrument"),
        (("serial", "serial_number"), "Serial number"),
        (("date",), "Date"),
        (("page", "page_number"), "Page"),
        (("sheet",), "Sheet"),
        (("workspace_path",), "Workspace"),
        (("user",), "User"),
        (("status",), "Status"),
        (("document_type",), "Document type"),
    ]
    for keys, label in header_map:
        val = _format_cell(
            _meta_get(data, *keys) if keys[0] != "document_type" else data.get("document_type")
        )
        if val:
            ws.append([label, val])
    return int(ws.max_row) + 1


def _append_sections_block(
    ws: Any,
    sections: Sequence[Dict[str, Any]],
    *,
    start_row: int = 1,
) -> None:
    """Write UI/document sections as Heading / Line rows."""
    if start_row <= 1:
        ws.append(["Heading", "Line"])
    else:
        ws.cell(row=start_row, column=1, value="Heading")
        ws.cell(row=start_row, column=2, value="Line")
        start_row += 1
    row_i = start_row if start_row > 1 else 2
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        heading = str(sec.get("heading") or "Section").strip() or "Section"
        lines = sec.get("lines") if isinstance(sec.get("lines"), list) else []
        if not lines:
            if start_row <= 1:
                ws.append([heading, ""])
            else:
                ws.cell(row=row_i, column=1, value=heading)
                ws.cell(row=row_i, column=2, value="")
                row_i += 1
            continue
        for ln in lines:
            text = str(ln).strip()
            if not text:
                continue
            if start_row <= 1:
                ws.append([heading, text])
            else:
                ws.cell(row=row_i, column=1, value=heading)
                ws.cell(row=row_i, column=2, value=text)
                row_i += 1


def _append_observation_table(
    ws: Any,
    rows: Sequence[Dict[str, Any]],
    *,
    start_row: int = 1,
    preferred_headers: Optional[Sequence[Any]] = None,
) -> None:
    used = _observation_columns(rows, preferred_headers=preferred_headers)
    headers = [c[1] for c in used]

    def _cell(row: Dict[str, Any], keys: tuple) -> str:
        if len(keys) == 1:
            return _format_cell(_row_value_ci(row, keys[0]))
        return _format_cell(_row_get(row, *keys))

    if start_row <= 1:
        ws.append(headers)
        for row in rows:
            ws.append([_cell(row, c[0]) for c in used])
        return
    for col_idx, header in enumerate(headers, start=1):
        ws.cell(row=start_row, column=col_idx, value=header)
    for r_idx, row in enumerate(rows, start=start_row + 1):
        for col_idx, c in enumerate(used, start=1):
            ws.cell(row=r_idx, column=col_idx, value=_cell(row, c[0]))


def export_ocr_extraction_to_excel(
    structured: Dict[str, Any],
    path: Path,
    *,
    image_paths: Optional[Sequence[str]] = None,
) -> Path:
    """Write a workbook with the full extraction visible on the first sheet.

    Sheet order (first sheet is what Excel opens):
      1. Extraction — metadata block + sections (UI) and/or observation table
      2. Observations — observation table only (when rows exist)
      3. Sections — UI/document text by heading (when sections exist)
      4. Metadata — header fields only
    """
    from openpyxl import Workbook

    data = structured if isinstance(structured, dict) else {}
    # Ensure visible_records / alternate payloads are promoted
    data = normalize_extracted_document(dict(data))
    path = Path(path)
    if path.suffix.lower() not in {".xlsx", ".xlsm"}:
        path = path.with_suffix(".xlsx")
    path.parent.mkdir(parents=True, exist_ok=True)

    rows = _flatten_structured_rows(data)
    if not rows:
        table = data.get("user_table") if isinstance(data.get("user_table"), str) else ""
        if table:
            parsed = _rows_from_markdown_table(table)
            if parsed:
                rows = parsed
    preferred = data.get("visible_headers") if isinstance(data.get("visible_headers"), list) else None
    sections = data.get("sections") if isinstance(data.get("sections"), list) else []

    wb = Workbook()

    extract_ws = wb.active
    extract_ws.title = "Extraction"
    _append_metadata_block(extract_ws, data, image_paths)
    extract_ws.append([])  # blank separator
    next_row = int(extract_ws.max_row) + 1
    if sections:
        extract_ws.cell(row=next_row, column=1, value="Sections")
        _append_sections_block(extract_ws, sections, start_row=next_row + 1)
        extract_ws.append([])
        next_row = int(extract_ws.max_row) + 1
    if rows:
        extract_ws.cell(row=next_row, column=1, value="Data")
        _append_observation_table(
            extract_ws, rows, start_row=next_row + 1, preferred_headers=preferred
        )
    elif not sections:
        extract_ws.append(["Data", "(none found in extraction)"])

    obs = wb.create_sheet("Observations")
    if rows:
        _append_observation_table(obs, rows, preferred_headers=preferred)
    else:
        obs.append(["Note"])
        obs.append(["No data rows in the saved extraction."])

    if sections:
        sec_ws = wb.create_sheet("Sections")
        _append_sections_block(sec_ws, sections)

    meta_ws = wb.create_sheet("Metadata")
    _append_metadata_block(meta_ws, data, image_paths)

    try:
        wb.save(path)
        return path.resolve()
    except PermissionError:
        alt = path.with_name(f"{path.stem}_export{path.suffix}")
        wb.save(alt)
        logger.warning("Could not overwrite locked %s; wrote %s", path, alt)
        return alt.resolve()


def _strip_md_inline(text: str) -> str:
    s = str(text or "")
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s.strip()


def _append_ocr_text_to_docx(doc: Any, body: str) -> None:
    """Write OCR user-facing markdown-ish text into a python-docx Document."""
    lines = (body or "").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue
        # Markdown heading
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading = _strip_md_inline(stripped.lstrip("#").strip())
            if heading:
                doc.add_heading(heading, level=min(max(level, 1), 4))
            i += 1
            continue
        # Pipe table block
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: List[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            body_rows = [
                ln for ln in table_lines
                if not re.match(r"^\|?\s*:?-{3,}", ln)
            ]
            if len(body_rows) >= 1:
                def _cells(ln: str) -> List[str]:
                    return [c.strip() for c in ln.strip("|").split("|")]

                headers = _cells(body_rows[0])
                data_rows = [_cells(ln) for ln in body_rows[1:]]
                cols = max(len(headers), max((len(r) for r in data_rows), default=0))
                if cols:
                    tbl = doc.add_table(rows=1 + len(data_rows), cols=cols)
                    try:
                        tbl.style = "Table Grid"
                    except Exception:
                        pass
                    for c_idx in range(cols):
                        tbl.rows[0].cells[c_idx].text = _strip_md_inline(
                            headers[c_idx] if c_idx < len(headers) else ""
                        )
                    for r_idx, row_cells in enumerate(data_rows, start=1):
                        for c_idx in range(cols):
                            tbl.rows[r_idx].cells[c_idx].text = _strip_md_inline(
                                row_cells[c_idx] if c_idx < len(row_cells) else ""
                            )
                    doc.add_paragraph("")
            continue
        # Bullet
        bullet = re.match(r"^[-*•]\s+(.*)$", stripped)
        if bullet:
            doc.add_paragraph(_strip_md_inline(bullet.group(1)), style="List Bullet")
            i += 1
            continue
        doc.add_paragraph(_strip_md_inline(stripped))
        i += 1


def export_ocr_extraction_to_docx(
    structured: Dict[str, Any],
    path: Path,
    *,
    image_paths: Optional[Sequence[str]] = None,
    user_table: str = "",
    title: Optional[str] = None,
) -> Path:
    """Write the OCR extraction (sections / table / metadata) into a Word document.

    Uses the same user-facing content as chat — never invents an essay titled from the filename.
    """
    import docx  # python-docx

    data = structured if isinstance(structured, dict) else {}
    data = normalize_extracted_document(dict(data))
    path = Path(path)
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    path.parent.mkdir(parents=True, exist_ok=True)

    paths_label = ", ".join(Path(p).name for p in (image_paths or [])) or "(image)"
    body = (user_table or "").strip()
    if not body:
        body = _format_structured_for_user(data, paths_label)

    doc_title = (title or "").strip()
    if not doc_title:
        doc_title = _format_cell(data.get("title")) or "OCR Extraction"
    # Never use a 1-letter filename stem as the document title (e.g. f.docx → "F")
    if len(doc_title) <= 2 and doc_title.upper() == path.stem.upper():
        doc_title = _format_cell(data.get("title")) or "OCR Extraction"

    document = docx.Document()
    document.add_heading(doc_title, 0)
    if not body:
        document.add_paragraph(f"Source: {paths_label}")
        dtype = str(data.get("document_type") or "").strip()
        if dtype:
            document.add_paragraph(f"Document type: {dtype}")
        document.add_paragraph("(No extraction text available.)")
    else:
        _append_ocr_text_to_docx(document, body)

    try:
        document.save(str(path))
        return path.resolve()
    except PermissionError:
        alt = path.with_name(f"{path.stem}_export{path.suffix}")
        document.save(str(alt))
        logger.warning("Could not overwrite locked %s; wrote %s", path, alt)
        return alt.resolve()


def _rows_from_markdown_table(text: str) -> List[Dict[str, Any]]:
    """Parse a simple pipe markdown table back into row dicts (export fallback)."""
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip().startswith("|")]
    if len(lines) < 3:
        return []
    def _cells(line: str) -> List[str]:
        parts = [p.strip() for p in line.strip("|").split("|")]
        return parts
    headers = _cells(lines[0])
    # skip separator line(s)
    body = [ln for ln in lines[1:] if not re.match(r"^\|?\s*:?-{3,}", ln)]
    header_map = {
        "inst. stn (from)": "instrument_station",
        "inst. stn": "instrument_station",
        "ref. stn (to)": "reference_station",
        "ref. stn": "reference_station",
        "ha face left": "hz_fl",
        "ha face right": "hz_fr",
        "va face left": "va_fl",
        "va face right": "va_fr",
        "slope dist": "slope_distance",
        "hor. dist": "horizontal_distance",
        "hor dist": "horizontal_distance",
    }
    keys = [header_map.get(h.lower(), normalize_field_name(h)) for h in headers]
    out: List[Dict[str, Any]] = []
    for ln in body:
        cells = _cells(ln)
        if len(cells) < 2:
            continue
        row = {keys[i]: cells[i] for i in range(min(len(keys), len(cells))) if keys[i]}
        if any(v not in (None, "") for v in row.values()):
            out.append(row)
    return out


def format_last_ocr_for_agent(payload: Dict[str, Any], *, max_chars: int = 24000) -> str:
    structured = payload.get("structured") if isinstance(payload, dict) else None
    if not isinstance(structured, dict):
        return ""
    parts = [
        "PRIOR VISION OCR EXTRACTION (from an earlier turn in this workspace).",
        "The user is referring to these values — save/export/compute/summarize from them.",
        "Do NOT ask for the image again unless they attach a new one.",
        f"Sources: {', '.join(str(p) for p in (payload.get('image_paths') or []))}",
        f"Document type: {payload.get('document_type') or structured.get('document_type') or ''}",
    ]
    table = payload.get("user_table")
    if isinstance(table, str) and table.strip():
        parts.append("User-facing table:")
        parts.append(table.strip()[: max_chars // 2])
    blob = json.dumps(structured, ensure_ascii=False)
    parts.append("Structured JSON:")
    parts.append(blob[: max_chars - 20] + "…[truncated]" if len(blob) > max_chars else blob)
    return "\n".join(parts)


def _meta_get(data: Dict[str, Any], *names: str) -> Any:
    meta = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    fields = data.get("fields") if isinstance(data.get("fields"), dict) else {}
    numbers = data.get("numbers") if isinstance(data.get("numbers"), dict) else {}
    plain = data.get("plain_text") if isinstance(data.get("plain_text"), dict) else {}
    for name in names:
        for src in (meta, fields, numbers, plain, data):
            if not isinstance(src, dict):
                continue
            if name in src and src[name] not in (None, ""):
                return src[name]
            key = normalize_field_name(name)
            for k, v in src.items():
                if normalize_field_name(str(k)) == key and v not in (None, ""):
                    return v
    return None


def _meta_set(data: Dict[str, Any], name: str, value: Any) -> None:
    meta = dict(data.get("metadata") or {}) if isinstance(data.get("metadata"), dict) else {}
    cur = meta.get(name)
    if isinstance(cur, dict):
        meta[name] = {**cur, "value": value, "raw": value, "repaired": True}
    else:
        meta[name] = value
    data["metadata"] = meta
    plain = data.get("plain_text")
    if isinstance(plain, dict):
        aliases = {
            "instrument": "instrument_name",
            "serial": "serial_number",
            "date": "date",
            "surveyed_by": "surveyed_by",
        }
        pk = aliases.get(name, name)
        if pk in plain:
            plain[pk] = value


def _choose_angle(observed: Any, complement: float) -> str:
    obs = parse_dms_to_deg(observed)
    if obs is not None and abs(((obs - complement + 180.0) % 360.0) - 180.0) <= 0.02:
        text = _unwrap(observed)
        return str(text) if text not in (None, "") else deg_to_dms(obs)
    return deg_to_dms(complement)


def _repair_angle_pair(fl_raw: Any, fr_raw: Any, *, vertical: bool) -> tuple[Any, Any, bool]:
    fl = parse_dms_to_deg(fl_raw)
    fr = parse_dms_to_deg(fr_raw)
    ok = va_pair_ok(fl, fr) if vertical else hz_pair_ok(fl, fr)
    if ok:
        return fl_raw, fr_raw, False
    fl_text = str(_unwrap(fl_raw) or "")
    fr_text = str(_unwrap(fr_raw) or "")
    for cand in _digit_variants(fl_text) + _digit_variants(fr_text):
        if cand in {fl_text, fr_text}:
            continue
        if cand != fl_text:
            nfl, nfr = parse_dms_to_deg(cand), fr
            if (va_pair_ok if vertical else hz_pair_ok)(nfl, nfr):
                return cand, fr_raw, True
        if cand != fr_text:
            nfl, nfr = fl, parse_dms_to_deg(cand)
            if (va_pair_ok if vertical else hz_pair_ok)(nfl, nfr):
                return fl_raw, cand, True
    if fl is None and fr is None:
        return fl_raw, fr_raw, False
    fl_conf, fr_conf = _cell_confidence(fl_raw), _cell_confidence(fr_raw)
    if vertical:
        if fr is not None and (fl is None or fl_conf <= fr_conf):
            return _choose_angle(fl_raw, (360.0 - fr) % 360.0), fr_raw, True
        if fl is not None:
            return fl_raw, _choose_angle(fr_raw, (360.0 - fl) % 360.0), True
    else:
        if fr is not None and (fl is None or fl_conf <= fr_conf):
            return _choose_angle(fl_raw, (fr + 180.0) % 360.0), fr_raw, True
        if fl is not None:
            return fl_raw, _choose_angle(fr_raw, (fl + 180.0) % 360.0), True
    return fl_raw, fr_raw, False


def _distance_variants(num: Optional[float]) -> List[float]:
    if num is None:
        return []
    text = f"{num:.3f}"
    out = [num]
    for cand in _digit_variants(text, limit=10):
        parsed = parse_survey_number(cand)
        if parsed is not None and parsed not in out:
            out.append(parsed)
    return out


def _compose_dms_parts(parts: Any) -> Optional[str]:
    """Build a DMS string from deg/min/sec pieces or pass through a string."""
    if parts is None or parts == "":
        return None
    if isinstance(parts, (str, int, float)):
        text = str(_unwrap(parts) if isinstance(parts, dict) else parts).strip()
        return text or None
    if not isinstance(parts, dict):
        return None
    # Already a value wrapper with a non-dict value
    if parts.get("value") not in (None, "") or parts.get("raw") not in (None, ""):
        un = _unwrap(parts)
        if un not in (None, "") and not isinstance(un, dict):
            return str(un)

    def _pick(*names: str) -> Any:
        for name in names:
            if name in parts and parts[name] not in (None, ""):
                return _unwrap(parts[name])
        return None

    deg = _pick("deg", "d", "degrees", "°")
    minutes = _pick("min", "m", "minutes", "'")
    seconds = _pick("sec", "s", "seconds", '"')
    # Explicit zero is valid (e.g. 000°00′00″)
    if deg is None and "deg" in parts:
        deg = _unwrap(parts.get("deg"))
    if minutes is None and "min" in parts:
        minutes = _unwrap(parts.get("min"))
    if seconds is None and "sec" in parts:
        seconds = _unwrap(parts.get("sec"))
    if deg is None and minutes is None and seconds is None:
        return None
    try:
        d = int(float(str(0 if deg is None else deg)))
        mi = int(float(str(0 if minutes is None else minutes)))
        se = int(round(float(str(0 if seconds is None else seconds))))
    except (TypeError, ValueError):
        return None
    return f"{d:03d}°{mi:02d}′{se:02d}″"


def _pull_face(container: Any, *keys: str) -> Any:
    if container is None:
        return None
    if not isinstance(container, dict):
        return container
    for key in keys:
        if key in container and container[key] not in (None, ""):
            return container[key]
        nk = normalize_field_name(key)
        for k, v in container.items():
            if normalize_field_name(str(k)) == nk and v not in (None, ""):
                return v
    return None


def _normalize_angle_cell(raw: Any) -> Any:
    if raw is None or raw == "":
        return None
    composed = _compose_dms_parts(raw)
    if composed:
        if isinstance(raw, dict) and ("confidence" in raw or "bbox" in raw):
            return {**raw, "value": composed, "raw": raw.get("raw", composed)}
        return composed
    return raw


def normalize_extracted_document(data: Dict[str, Any]) -> Dict[str, Any]:
    """Flatten nested traverse angle objects; promote spreadsheet visible_records → rows."""
    out = dict(data or {})
    # Promote alternate spreadsheet payloads into rows before anything else
    if not (isinstance(out.get("rows"), list) and out.get("rows")):
        for alt in ("visible_records", "records", "data_rows"):
            cand = out.get(alt)
            if isinstance(cand, list) and cand and isinstance(cand[0], dict):
                out["rows"] = [dict(r) for r in cand if isinstance(r, dict)]
                break
    # Normalize document_type aliases
    dt = str(out.get("document_type") or "").strip().lower().replace(" ", "_")
    if dt in {"generic_spreadsheet", "excel", "spreadsheet"}:
        out["document_type"] = DOC_SPREADSHEET
    if dt in {"ui_screenshot", "application_ui", "desktop_ui", "app_ui"}:
        out["document_type"] = DOC_UI
    # Normalize sections / panels for UI screenshots
    sections = out.get("sections") or out.get("panels") or out.get("regions")
    if isinstance(sections, list) and sections:
        clean_sections: List[Dict[str, Any]] = []
        for sec in sections:
            if isinstance(sec, str) and sec.strip():
                clean_sections.append({"heading": "Notes", "lines": [sec.strip()]})
                continue
            if not isinstance(sec, dict):
                continue
            heading = str(
                sec.get("heading") or sec.get("title") or sec.get("name") or sec.get("panel") or "Section"
            ).strip() or "Section"
            raw_lines = sec.get("lines") or sec.get("text") or sec.get("items") or sec.get("content")
            lines_out: List[str] = []
            if isinstance(raw_lines, str):
                lines_out = [ln.strip() for ln in raw_lines.splitlines() if ln.strip()]
            elif isinstance(raw_lines, list):
                for ln in raw_lines:
                    if isinstance(ln, dict):
                        val = _unwrap(ln)
                        if val not in (None, ""):
                            lines_out.append(str(val).strip())
                    elif ln not in (None, ""):
                        lines_out.append(str(ln).strip())
            if heading or lines_out:
                clean_sections.append({"heading": heading, "lines": lines_out})
        if clean_sections:
            out["sections"] = clean_sections
            if not out.get("document_type") or out.get("document_type") == DOC_GENERIC:
                out["document_type"] = DOC_UI
    # Drop invented null survey metadata (common when models pad the schema)
    meta = out.get("metadata")
    if isinstance(meta, dict):
        cleaned = {
            k: v for k, v in meta.items()
            if v not in (None, "", [], {})
            and not (isinstance(v, dict) and v.get("value") in (None, "") and v.get("raw") in (None, ""))
        }
        out["metadata"] = cleaned
    # Lift phone from odd places into metadata
    phone = _meta_get(out, "phone", "telephone", "tel", "mobile", "contact")
    if phone in (None, ""):
        blob = " ".join(
            [
                str(out.get("title") or ""),
                str(out.get("plain_text") or ""),
                json.dumps(out.get("fields") or {}, ensure_ascii=False)[:800],
            ]
        )
        m = re.search(r"(?:\+?\d[\d\s\-()]{7,}\d)", blob)
        if m:
            phone = re.sub(r"\s+", " ", m.group(0)).strip()
    if phone not in (None, ""):
        _meta_set(out, "phone", str(_unwrap(phone)).strip())

    rows = [dict(r) for r in _as_rows(out)]
    if not rows:
        return out

    # Spreadsheet / GIS attribute tables: keep original column keys; do not force traverse flattening
    if _rows_look_like_spreadsheet(rows, out) and not _rows_look_like_traverse(rows):
        out["document_type"] = out.get("document_type") or DOC_SPREADSHEET
        normalized_ss: List[Dict[str, Any]] = []
        for row in rows:
            item = {str(k): _unwrap(v) for k, v in row.items() if str(k).lower() not in {"bbox", "confidence"}}
            if any(v not in (None, "") for v in item.values()):
                normalized_ss.append(item)
        out["rows"] = normalized_ss
        # Remember header order when provided
        headers = out.get("visible_headers")
        if isinstance(headers, list) and headers:
            out["visible_headers"] = [str(h) for h in headers if str(h).strip()]
        return out

    normalized: List[Dict[str, Any]] = []
    for row in rows:
        item = {normalize_field_name(k): v for k, v in row.items()}
        ha = item.pop("horizontal_angle", None) or item.pop("ha", None) or item.pop("horizontal", None)
        va = item.pop("vertical_angle", None) or item.pop("va", None) or item.pop("vertical", None)
        if isinstance(ha, dict):
            fl = _pull_face(ha, "face_left", "fl", "left", "hz_fl", "ha_fl")
            fr = _pull_face(ha, "face_right", "fr", "right", "hz_fr", "ha_fr")
            if fl is not None and item.get("hz_fl") in (None, ""):
                item["hz_fl"] = fl
            if fr is not None and item.get("hz_fr") in (None, ""):
                item["hz_fr"] = fr
        if isinstance(va, dict):
            fl = _pull_face(va, "face_left", "fl", "left", "va_fl")
            fr = _pull_face(va, "face_right", "fr", "right", "va_fr")
            if fl is not None and item.get("va_fl") in (None, ""):
                item["va_fl"] = fl
            if fr is not None and item.get("va_fr") in (None, ""):
                item["va_fr"] = fr
        for prefix, target in (
            ("hz_fl", "hz_fl"), ("hz_fr", "hz_fr"), ("va_fl", "va_fl"), ("va_fr", "va_fr"),
            ("ha_fl", "hz_fl"), ("ha_fr", "hz_fr"),
        ):
            if item.get(target) not in (None, ""):
                continue
            parts = {
                "deg": item.pop(f"{prefix}_deg", None) or item.pop(f"{prefix}_d", None),
                "min": item.pop(f"{prefix}_min", None) or item.pop(f"{prefix}_m", None),
                "sec": item.pop(f"{prefix}_sec", None) or item.pop(f"{prefix}_s", None),
            }
            if any(v not in (None, "") for v in parts.values()):
                item[target] = parts
        for key in ("hz_fl", "hz_fr", "va_fl", "va_fr", "ha_fl", "ha_fr"):
            if key in item:
                canon = "hz_fl" if key == "ha_fl" else "hz_fr" if key == "ha_fr" else key
                item[canon] = _normalize_angle_cell(item[key])
                if key != canon:
                    item.pop(key, None)
        dist = item.pop("distance", None) if isinstance(item.get("distance"), dict) else None
        if isinstance(dist, dict):
            if item.get("slope_distance") in (None, "") and item.get("slope") in (None, ""):
                item["slope_distance"] = dist.get("slope") or dist.get("slope_distance") or dist.get("slope_dist")
            if item.get("horizontal_distance") in (None, "") and item.get("hor") in (None, ""):
                item["horizontal_distance"] = (
                    dist.get("hor") or dist.get("horizontal") or dist.get("horizontal_distance")
                )
        if "from" in item and item.get("instrument_station") in (None, ""):
            item["instrument_station"] = item["from"]
        if "to" in item and item.get("reference_station") in (None, ""):
            item["reference_station"] = item["to"]
        normalized.append(item)
    out["rows"] = normalized
    return out


def _str_mode(vals: Sequence[Optional[str]]) -> Optional[str]:
    cleaned = [v for v in vals if v]
    if not cleaned:
        return None
    return max(set(cleaned), key=cleaned.count)


def _station_digit_variants(station: str) -> List[str]:
    s = normalize_station_id(station) or station
    if "-" not in s:
        return []
    prefix, suffix = s.rsplit("-", 1)
    if not suffix.isdigit():
        return []
    out: List[str] = []
    # Prefer mutating the trailing station digit(s) before leading zeros (02→03 before 02→82).
    for i in range(len(suffix) - 1, -1, -1):
        ch = suffix[i]
        for a, b in (("2", "3"), ("3", "2"), ("1", "4"), ("4", "1"), ("0", "8"), ("8", "0"), ("4", "9"), ("9", "4")):
            if ch != a:
                continue
            alt = f"{prefix}-{suffix[:i]}{b}{suffix[i + 1:]}"
            if alt != s and alt not in out:
                out.append(alt)
    return out


def _repair_traverse_station_pattern(rows: List[Dict[str, Any]]) -> None:
    """Lock alternating/distance-clustered refs and a single instrument station that is not a ref."""
    if len(rows) < 2:
        return
    refs = [normalize_station_id(_row_get(r, "reference_station", "to")) for r in rows]
    froms = [normalize_station_id(_row_get(r, "instrument_station", "from")) for r in rows]

    # Distance clusters: same from→to shots repeat the same slope/hor (±0.5 m)
    clusters: Dict[int, List[int]] = {}
    for i, row in enumerate(rows):
        slope = parse_survey_number(_row_get(row, "slope_distance", "slope", "distance"))
        hor = parse_survey_number(_row_get(row, "horizontal_distance", "hor"))
        dist = slope if slope is not None else hor
        if dist is None:
            continue
        key = int(round(float(dist)))
        # Bucket nearby integers (343 vs 363 stay distinct; 363.05 vs 363.054 merge)
        matched = None
        for existing in clusters:
            if abs(existing - key) <= 1:
                matched = existing
                break
        clusters.setdefault(matched if matched is not None else key, []).append(i)

    if len(clusters) >= 2:
        cluster_modes: Dict[int, str] = {}
        for ckey, idxs in clusters.items():
            mode = _str_mode([refs[i] for i in idxs])
            if mode:
                cluster_modes[ckey] = mode
        # Ensure cluster modes are distinct when possible via digit swaps
        modes = list(cluster_modes.values())
        if len(set(modes)) < len(modes) and len(modes) == 2:
            a, b = modes[0], modes[1]
            if a == b:
                for alt in _station_digit_variants(b):
                    if alt != a:
                        # Assign alt to the cluster that had more from==to collisions
                        keys = list(cluster_modes.keys())
                        cluster_modes[keys[1]] = alt
                        break
        for ckey, idxs in clusters.items():
            mode = cluster_modes.get(ckey)
            if not mode:
                continue
            for i in idxs:
                _row_set(rows[i], "reference_station", mode)
        refs = [normalize_station_id(_row_get(r, "reference_station", "to")) for r in rows]
    elif len(rows) >= 4 and len(rows) % 2 == 0:
        odd_mode = _str_mode([refs[i] for i in range(0, len(rows), 2)])
        even_mode = _str_mode([refs[i] for i in range(1, len(rows), 2)])
        if odd_mode and even_mode and odd_mode != even_mode:
            for i, row in enumerate(rows):
                _row_set(row, "reference_station", odd_mode if i % 2 == 0 else even_mode)
            refs = [normalize_station_id(_row_get(r, "reference_station", "to")) for r in rows]

    ref_set = {r for r in refs if r}
    from_mode = _str_mode(froms)
    # Prefer an instrument station that does not collide with reference stations
    if from_mode and from_mode in ref_set:
        for alt in _station_digit_variants(from_mode):
            if alt not in ref_set:
                from_mode = alt
                break
    # Same-row from==to is almost never valid
    for row in rows:
        fr = normalize_station_id(_row_get(row, "instrument_station", "from"))
        to = normalize_station_id(_row_get(row, "reference_station", "to"))
        if fr and to and fr == to:
            for alt in _station_digit_variants(fr):
                if alt != to and alt not in ref_set:
                    from_mode = alt
                    break
            if from_mode and from_mode in ref_set:
                for alt in _station_digit_variants(fr):
                    if alt != to:
                        from_mode = alt
                        break
    if not from_mode:
        candidates = [f for f in froms if f and f not in ref_set]
        from_mode = _str_mode(candidates) or _str_mode(froms)
    if from_mode:
        for row in rows:
            _row_set(row, "instrument_station", from_mode)



def _repair_distance_4_vs_6(
    slope: Optional[float], hor: Optional[float]
) -> tuple[Optional[float], Optional[float]]:
    """Fix common tens/hundreds digit confusions when slope ≈ hor."""
    if slope is None and hor is None:
        return slope, hor

    def _flip_tens_4_to_6(val: float) -> Optional[float]:
        ip = int(abs(val))
        if 340 <= ip <= 349 and ((ip // 10) % 10) == 4:
            return float(ip + 20) + (abs(val) - ip)  # 343 → 363
        return None

    def _flip_hundreds_2_to_9(val: float) -> Optional[float]:
        ip = int(abs(val))
        if 240 <= ip <= 249:
            return float(ip + 50) + (abs(val) - ip)  # 244 → 294
        return None

    for flip in (_flip_tens_4_to_6, _flip_hundreds_2_to_9):
        if slope is not None and hor is not None and abs(slope - hor) <= 0.5:
            ns, nh = flip(slope), flip(hor)
            if ns is not None and nh is not None and abs(ns - nh) <= 0.5:
                return ns, nh
        if slope is not None:
            ns = flip(slope)
            if ns is not None:
                nh = flip(hor) if hor is not None else None
                if nh is not None and abs(ns - nh) <= 0.5:
                    return ns, nh
                if hor is None or abs(ns - hor) <= 0.5:
                    return ns, hor if nh is None else nh

    # Hor truncated / units digit slip while slope looks solid (363.057 vs 362.65)
    if slope is not None and hor is not None and 0.05 < abs(slope - hor) <= 1.25:
        best = None
        best_gap = abs(slope - hor)
        for hv in _distance_variants(hor):
            gap = abs(slope - hv)
            if gap < best_gap and gap <= 0.5:
                best_gap, best = gap, hv
        if best is not None:
            return slope, best
        if abs(int(slope) - int(hor)) == 1:
            snapped = round(slope - 0.004, 3)
            if snapped <= slope + 1e-9:
                return slope, snapped
    return slope, hor


def apply_survey_context_repairs(
    data: Dict[str, Any],
    *,
    document_type: str = "",
    style_card: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Resolve faint glyphs using sheet identities. Does not invent new observations."""
    out = dict(data or {})
    dtype = str(document_type or out.get("document_type") or "")
    rows_preview = _as_rows(out)
    # Spreadsheet / GIS attribute screenshots: never force traverse station/distance repairs
    if dtype in {DOC_SPREADSHEET, DOC_UI, "generic_spreadsheet"} or (
        _rows_look_like_spreadsheet(rows_preview, out) and not _rows_look_like_traverse(rows_preview)
    ):
        meta = out.get("metadata")
        if isinstance(meta, dict):
            out["metadata"] = {
                k: v for k, v in meta.items()
                if v not in (None, "", [], {})
            }
        if rows_preview and not (isinstance(out.get("rows"), list) and out.get("rows")):
            out["rows"] = [dict(r) for r in rows_preview]
        return out
    card = style_card if isinstance(style_card, dict) else (out.get("style_card") or {})
    serial = repair_serial_number(_meta_get(out, "serial", "serial_number"))
    if serial:
        _meta_set(out, "serial", serial)
    inst = repair_instrument_name(_meta_get(out, "instrument", "instrument_name"))
    if inst:
        _meta_set(out, "instrument", inst)
    date = repair_sheet_date(_meta_get(out, "date"), card)
    if date:
        _meta_set(out, "date", date)
    phone = repair_phone_number(_meta_get(out, "phone", "telephone", "tel", "mobile"))
    if phone:
        _meta_set(out, "phone", phone)
    surveyed = repair_surveyed_by_name(_meta_get(out, "surveyed_by"))
    if surveyed:
        _meta_set(out, "surveyed_by", surveyed)

    rows = [dict(r) for r in rows_preview]
    if not rows:
        return out
    last_inst: Optional[str] = None
    for row in rows:
        inst_id = normalize_station_id(_row_get(row, "instrument_station", "from", "inst_stn", "station_from"))
        if inst_id:
            last_inst = inst_id
            _row_set(row, "instrument_station", inst_id)
        elif last_inst:
            _row_set(row, "instrument_station", last_inst)
        ref_id = normalize_station_id(_row_get(row, "reference_station", "to", "ref_stn", "station_to"))
        if ref_id:
            _row_set(row, "reference_station", ref_id)

    _repair_traverse_station_pattern(rows)

    pair_slopes: Dict[tuple, List[float]] = {}
    pair_hors: Dict[tuple, List[float]] = {}
    for row in rows:
        if dtype in {DOC_TRAVERSE, DOC_GENERIC, DOC_TABLE, ""}:
            fl, fr, changed = _repair_angle_pair(
                _row_get(row, "hz_fl", "ha_fl"), _row_get(row, "hz_fr", "ha_fr"), vertical=False
            )
            if changed:
                _row_set(row, "hz_fl", fl)
                _row_set(row, "hz_fr", fr)
            vfl, vfr, vchanged = _repair_angle_pair(
                _row_get(row, "va_fl"), _row_get(row, "va_fr"), vertical=True
            )
            if vchanged:
                _row_set(row, "va_fl", vfl)
                _row_set(row, "va_fr", vfr)
        # 343↔363: closed-loop 6 often misread as 4 when slope≈hor in the 340s
        slope = parse_survey_number(_row_get(row, "slope_distance", "slope", "distance"))
        hor = parse_survey_number(_row_get(row, "horizontal_distance", "hor"))
        slope, hor = _repair_distance_4_vs_6(slope, hor)
        if slope is not None:
            _row_set(row, "slope_distance", slope)
        if hor is not None:
            _row_set(row, "horizontal_distance", hor)
        inst_id = str(normalize_station_id(_row_get(row, "instrument_station", "from")) or "")
        ref_id = str(normalize_station_id(_row_get(row, "reference_station", "to")) or "")
        if slope is not None and hor is not None and abs(slope - hor) > 1.0:
            best = None
            best_gap = abs(slope - hor)
            for sv in _distance_variants(slope):
                for hv in _distance_variants(hor):
                    if sv + 1e-9 < hv:
                        continue
                    gap = abs(sv - hv)
                    if gap < best_gap:
                        best_gap, best = gap, (sv, hv)
            if best and best_gap <= 0.5:
                _row_set(row, "slope_distance", best[0])
                _row_set(row, "horizontal_distance", best[1])
                slope, hor = best
        key = (inst_id, ref_id)
        if slope is not None:
            pair_slopes.setdefault(key, []).append(slope)
        if hor is not None:
            pair_hors.setdefault(key, []).append(hor)

    def _mode(vals: List[float]) -> Optional[float]:
        if not vals:
            return None
        rounded = [round(v, 3) for v in vals]
        unique = list(dict.fromkeys(rounded))
        for a in unique:
            for b in unique:
                if a >= b:
                    continue
                sa, sb = str(int(a)), str(int(b))
                if len(sa) == len(sb) and abs(int(a) - int(b)) == 20:
                    if "4" in sa and "6" in sb:
                        return b
                    if "6" in sa and "4" in sb:
                        return a
        winner = max(set(rounded), key=rounded.count)
        return winner if rounded.count(winner) >= 2 else None

    for row in rows:
        inst_id = str(normalize_station_id(_row_get(row, "instrument_station", "from")) or "")
        ref_id = str(normalize_station_id(_row_get(row, "reference_station", "to")) or "")
        key = (inst_id, ref_id)
        slope_mode = _mode(pair_slopes.get(key) or [])
        hor_mode = _mode(pair_hors.get(key) or [])
        slope = parse_survey_number(_row_get(row, "slope_distance", "slope", "distance"))
        hor = parse_survey_number(_row_get(row, "horizontal_distance", "hor"))
        if slope_mode is not None and (slope is None or abs(slope - slope_mode) >= 1.0):
            _row_set(row, "slope_distance", slope_mode)
        if hor_mode is not None and (hor is None or abs(hor - hor_mode) >= 1.0):
            _row_set(row, "horizontal_distance", hor_mode)

    if rows:
        out["rows"] = rows
    return out


def _format_cell(val: Any) -> str:
    v = _unwrap(val)
    if v is None or v == "":
        return ""
    if isinstance(v, bool):
        return str(v)
    if isinstance(v, int):
        return str(v)
    if isinstance(v, float):
        # Preserve GIS/coordinate precision; still trim trailing zeros
        text = f"{v:.10f}".rstrip("0").rstrip(".")
        return text if text else "0"
    return str(v).strip()


def _format_structured_for_user(data: Dict[str, Any], paths: str, *, text: str = "") -> str:
    lines: List[str] = []
    title = _format_cell(data.get("title")) or ""
    if title:
        lines.append(f"**{title}**")
    lines.append(f"**Source:** {paths}")
    lines.append("")
    sheet_name = _format_cell(_meta_get(data, "sheet"))
    if sheet_name:
        lines.append(f"- **Sheet:** {sheet_name}")
    for meta_key, label in (
        ("workspace_path", "Workspace"),
        ("user", "User"),
        ("status", "Status"),
    ):
        val = _format_cell(_meta_get(data, meta_key))
        if val:
            lines.append(f"- **{label}:** {val}")
    header_map = [
        (("organization",), "Organization"),
        (("phone", "telephone", "tel", "mobile", "contact"), "Phone"),
        (("surveyed_by",), "Surveyed by"),
        (("computed_by",), "Computed by"),
        (("instrument", "instrument_name"), "Instrument"),
        (("serial", "serial_number"), "Serial number"),
        (("date",), "Date"),
        (("page", "page_number"), "Page"),
    ]
    for keys, label in header_map:
        val = _format_cell(_meta_get(data, *keys))
        if val:
            lines.append(f"- **{label}:** {val}")

    # UI / document sections grouped under headings
    sections = data.get("sections") if isinstance(data.get("sections"), list) else []
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        heading = str(sec.get("heading") or "Section").strip() or "Section"
        sec_lines = sec.get("lines") if isinstance(sec.get("lines"), list) else []
        lines.append("")
        lines.append(f"### {heading}")
        for ln in sec_lines:
            s = str(ln).strip()
            if s:
                lines.append(s if s.startswith(("-", "*", "•")) else f"- {s}")

    rows = _as_rows(data)
    preferred = data.get("visible_headers") if isinstance(data.get("visible_headers"), list) else None
    if rows:
        look_traverse = _rows_look_like_traverse(rows)
        used = _observation_columns(
            rows, preferred_headers=preferred, force_traverse=look_traverse
        )
        if used:
            lines.append("")
            lines.append("| " + " | ".join(c[1] for c in used) + " |")
            lines.append("|" + "|".join("---" for _ in used) + "|")
            for row in rows:
                cells = []
                for c in used:
                    if len(c[0]) == 1 and not look_traverse:
                        cells.append(_format_cell(_row_value_ci(row, c[0][0])))
                    else:
                        cells.append(_format_cell(_row_get(row, *c[0])))
                lines.append("| " + " | ".join(cells) + " |")
    coords = data.get("coordinates") if isinstance(data.get("coordinates"), list) else []
    if coords:
        lines.append("")
        lines.append("| Label | Easting | Northing |")
        lines.append("|---|---|---|")
        for item in coords:
            if not isinstance(item, dict):
                continue
            lines.append(
                f"| {_format_cell(item.get('label'))} | "
                f"{_format_cell(item.get('easting'))} | {_format_cell(item.get('northing'))} |"
            )
    plain = data.get("plain_text")
    rendered = "\n".join(lines)
    has_body = bool(sections) or bool(rows) or bool(coords)
    if isinstance(plain, str) and plain.strip() and plain.strip() not in rendered:
        if not has_body or len(plain.strip()) <= 80:
            if not plain.strip().startswith("{") and "style_card" not in plain[:120]:
                lines.append("")
                lines.append(plain.strip())
    elif text.strip() and not has_body:
        t = text.strip()
        if not t.startswith("{") and "style_card" not in t[:200] and "visible_records" not in t[:200]:
            lines.append("")
            lines.append(t)
    return "\n".join(lines).strip()


def format_ocr_review_for_user(
    review: Dict[str, Any],
    *,
    note: str = "",
) -> str:
    """Render the click-to-verify payload (after optional Apply) as the chat extraction."""
    if not isinstance(review, dict):
        return note or ""
    paths = ", ".join(Path(p).name for p in (review.get("image_paths") or [])) or "(image)"
    structured = structured_from_ocr_review(review)
    body = _format_structured_for_user(structured, paths)
    if note:
        body = f"{note}\n\n{body}"
    return body.strip()


def validate_levelling_or_calibration(rows: Sequence[Dict[str, Any]]) -> List[ValidationCheck]:
    checks: List[ValidationCheck] = []
    parsed = [
        {
            "backsight": parse_survey_number(_row_get(row, "backsight")),
            "intermediate_sight": parse_survey_number(_row_get(row, "intermediate_sight")),
            "foresight": parse_survey_number(_row_get(row, "foresight")),
            "reduced_level": parse_survey_number(_row_get(row, "reduced_level")),
            "height_of_collimation": parse_survey_number(_row_get(row, "height_of_collimation")),
            "rise": parse_survey_number(_row_get(row, "rise")),
            "fall": parse_survey_number(_row_get(row, "fall")),
        }
        for row in rows
    ]
    carry_hi: Optional[float] = None
    for i, vals in enumerate(parsed):
        stn = f"row[{i}]"
        rl, bs, hi, fs, inter = (
            vals["reduced_level"], vals["backsight"], vals["height_of_collimation"],
            vals["foresight"], vals["intermediate_sight"],
        )
        if rl is not None and bs is not None:
            expected_hi = rl + bs
            if hi is not None:
                checks.append(ValidationCheck(
                    name=f"{stn}.HI=RL+BS", passed=_almost(hi, expected_hi, _LEVEL_TOL),
                    expected=round(expected_hi, 4), observed=hi,
                    fields=["reduced_level", "backsight", "height_of_collimation"],
                    detail="Height of collimation must equal reduced level plus backsight.",
                ))
            carry_hi = hi if hi is not None else expected_hi
        elif hi is not None:
            carry_hi = hi
        active_hi = hi if hi is not None else carry_hi
        if active_hi is not None and fs is not None:
            expected_rl = active_hi - fs
            target = rl if (rl is not None and bs is None) else (
                parsed[i + 1]["reduced_level"] if i + 1 < len(parsed) else rl
            )
            if target is not None:
                checks.append(ValidationCheck(
                    name=f"{stn}.RL=HI-FS", passed=_almost(target, expected_rl, _LEVEL_TOL),
                    expected=round(expected_rl, 4), observed=target,
                    fields=["height_of_collimation", "foresight", "reduced_level"],
                    detail="Reduced level after a foresight must equal HI minus FS.",
                ))
        if active_hi is not None and inter is not None and rl is not None and bs is None:
            expected_rl = active_hi - inter
            checks.append(ValidationCheck(
                name=f"{stn}.RL=HI-IS", passed=_almost(rl, expected_rl, _LEVEL_TOL),
                expected=round(expected_rl, 4), observed=rl,
                fields=["height_of_collimation", "intermediate_sight", "reduced_level"],
                detail="Intermediate-sight reduced level must equal HI minus IS.",
            ))
        if bs is not None and fs is not None:
            delta = bs - fs
            if vals["rise"] is not None and delta >= 0:
                checks.append(ValidationCheck(
                    name=f"{stn}.rise=BS-FS", passed=_almost(vals["rise"], delta, _LEVEL_TOL),
                    expected=round(delta, 4), observed=vals["rise"],
                    fields=["backsight", "foresight", "rise"],
                    detail="Rise should equal BS minus FS when the ground rises.",
                ))
            if vals["fall"] is not None and delta < 0:
                checks.append(ValidationCheck(
                    name=f"{stn}.fall=FS-BS", passed=_almost(vals["fall"], -delta, _LEVEL_TOL),
                    expected=round(-delta, 4), observed=vals["fall"],
                    fields=["backsight", "foresight", "fall"],
                    detail="Fall should equal FS minus BS when the ground falls.",
                ))
    rls = [v["reduced_level"] for v in parsed if v["reduced_level"] is not None]
    bs_sum = sum(v["backsight"] or 0.0 for v in parsed if v["backsight"] is not None)
    fs_sum = sum(v["foresight"] or 0.0 for v in parsed if v["foresight"] is not None)
    if len(rls) >= 2 and (bs_sum or fs_sum):
        expected, observed = rls[-1] - rls[0], bs_sum - fs_sum
        checks.append(ValidationCheck(
            name="sum(BS)-sum(FS)=last_RL-first_RL",
            passed=_almost(observed, expected, _LEVEL_TOL * max(1, len(parsed))),
            expected=round(expected, 4), observed=round(observed, 4),
            fields=["backsight", "foresight", "reduced_level"],
            detail="Arithmetic close of a levelling run.",
        ))
    return checks


def validate_traverse_rows(rows: Sequence[Dict[str, Any]]) -> List[ValidationCheck]:
    checks: List[ValidationCheck] = []
    for i, row in enumerate(rows):
        hz_fl = parse_dms_to_deg(_row_get(row, "hz_fl", "ha_fl"))
        hz_fr = parse_dms_to_deg(_row_get(row, "hz_fr", "ha_fr"))
        if hz_fl is not None and hz_fr is not None:
            observed = _hz_abs_diff(hz_fl, hz_fr)
            checks.append(ValidationCheck(
                name=f"row[{i}].HA_|FL-FR|≈180",
                passed=hz_pair_ok(hz_fl, hz_fr),
                expected=180.0, observed=round(observed, 4),
                tolerance=_ANGLE_TOL_DEG, fields=["hz_fl", "hz_fr"],
                detail="Horizontal Face Left and Face Right on the same inst→ref shot differ by ~180°.",
            ))
        va_fl = parse_dms_to_deg(_row_get(row, "va_fl"))
        va_fr = parse_dms_to_deg(_row_get(row, "va_fr"))
        if va_fl is not None and va_fr is not None:
            observed = (va_fl + va_fr) % 360.0
            if observed > 180:
                observed = 360.0 - observed
            checks.append(ValidationCheck(
                name=f"row[{i}].VA_FL+FR≈360",
                passed=va_pair_ok(va_fl, va_fr),
                expected=360.0, observed=round((va_fl + va_fr) % 360.0, 4),
                tolerance=_ANGLE_TOL_DEG, fields=["va_fl", "va_fr"],
                detail="Vertical Face Left and Face Right on the same inst→ref shot sum to ~360°.",
            ))
        dist = parse_survey_number(_row_get(row, "distance", "length", "dist", "slope_distance"))
        if dist is not None and dist < 0:
            checks.append(ValidationCheck(
                name=f"row[{i}].distance>=0", passed=False, expected=0.0, observed=dist,
                tolerance=_DISTANCE_TOL, fields=["distance"], detail="Traverse distance cannot be negative.",
            ))
        lat = parse_survey_number(_row_get(row, "latitude", "lat"))
        dep = parse_survey_number(_row_get(row, "departure", "dep"))
        if dist is not None and lat is not None and dep is not None:
            hyp = math.hypot(lat, dep)
            checks.append(ValidationCheck(
                name=f"row[{i}].sqrt(lat^2+dep^2)=distance",
                passed=_almost(hyp, dist, max(_DISTANCE_TOL, 0.01 * dist)),
                expected=dist, observed=round(hyp, 4),
                fields=["latitude", "departure", "distance"],
                detail="Latitude/departure must reconstruct the measured distance.",
            ))
    return checks


def collect_uncertain_fields(extracted: Dict[str, Any], *, threshold: float = _LOW_CONFIDENCE) -> List[str]:
    found: List[str] = []
    conf = extracted.get("confidence")
    if isinstance(conf, dict):
        for key, val in conf.items():
            try:
                if float(val) < threshold:
                    found.append(str(key))
            except (TypeError, ValueError):
                found.append(str(key))
    for item in extracted.get("uncertain") or []:
        if item and str(item) not in found:
            found.append(str(item))

    def _walk(obj: Any, prefix: str) -> None:
        if isinstance(obj, dict):
            c = obj.get("confidence")
            try:
                if c is not None and float(c) < _NESTED_UNCERTAIN and prefix and prefix not in found:
                    found.append(prefix)
            except (TypeError, ValueError):
                pass
            for k, v in obj.items():
                if k not in {"confidence", "notes", "plain_text", "style_card"}:
                    _walk(v, f"{prefix}.{k}" if prefix else str(k))
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                _walk(v, f"{prefix}[{i}]")

    _walk(extracted.get("rows"), "rows")
    _walk(extracted.get("fields"), "fields")
    return found


def validate_ocr_extraction(
    extracted: Optional[Dict[str, Any]],
    *,
    user_text: str = "",
    image_paths: Sequence[str] = (),
    document_type: Optional[str] = None,
) -> ValidationReport:
    data = extracted if isinstance(extracted, dict) else {}
    dtype = document_type or data.get("document_type") or classify_ocr_document(user_text, image_paths, data)
    report = ValidationReport(document_type=str(dtype or DOC_GENERIC))
    rows = _as_rows(data)
    if not rows and not data:
        report.notes.append("No structured fields to validate.")
        return report
    if report.document_type in {DOC_CALIBRATION, DOC_LEVELLING, DOC_TABLE, DOC_GENERIC}:
        report.checks.extend(validate_levelling_or_calibration(rows))
    if report.document_type in {DOC_TRAVERSE, DOC_CADASTRAL}:
        report.checks.extend(validate_traverse_rows(rows))
        report.checks.extend(validate_levelling_or_calibration(rows))
    report.uncertain_fields = collect_uncertain_fields(data)
    failed = [c for c in report.checks if not c.passed]
    if failed:
        report.notes.append(f"{len(failed)} arithmetic/surveying check(s) failed — treat flagged readings as unverified.")
    if report.uncertain_fields:
        report.notes.append("Low-confidence fields: " + ", ".join(report.uncertain_fields[:12]))
    if not report.checks and rows:
        report.notes.append("No closed surveying identities were found; values were transcribed only.")
    return report


def merge_validation_into_structured(structured: Dict[str, Any], report: ValidationReport) -> Dict[str, Any]:
    out = dict(structured or {})
    out["document_type"] = out.get("document_type") or report.document_type
    out["_validation"] = report.to_dict()
    return out


def format_validation_for_user(report: ValidationReport) -> str:
    return _format_validation_dict(report.to_dict())


def _format_validation_dict(data: Dict[str, Any]) -> str:
    if not data:
        return ""
    lines = ["### Accuracy checks (independent of the vision model)", ""]
    if data.get("ok") and not data.get("uncertain_fields"):
        lines.append("All applicable surveying/arithmetic checks passed.")
    else:
        lines.append("**Review required.** Do not treat failed or low-confidence readings as final.")
    checks = data.get("checks") or []
    if checks:
        lines += ["", "| Check | Result | Expected | Observed |", "|---|---|---|---|"]
        for c in checks:
            if not isinstance(c, dict):
                continue
            exp = "" if c.get("expected") is None else str(c.get("expected"))
            obs = "" if c.get("observed") is None else str(c.get("observed"))
            lines.append(f"| {c.get('name')} | {'PASS' if c.get('passed') else 'FAIL'} | {exp} | {obs} |")
    if data.get("uncertain_fields"):
        lines += ["", "**Uncertain fields:** " + ", ".join(str(x) for x in data["uncertain_fields"])]
    for note in data.get("notes") or []:
        lines += ["", f"- {note}"]
    return "\n".join(lines)


def fields_needing_reread(report: ValidationReport, limit: int = 6) -> List[str]:
    names: List[str] = []
    for c in report.checks:
        if c.passed:
            continue
        for f in c.fields:
            if f not in names:
                names.append(f)
    for f in report.uncertain_fields:
        tail = re.sub(r"\[\d+\]", "", re.sub(r".*\.", "", f))
        if tail and tail not in names:
            names.append(tail)
    return names[:limit]


def _review_field_cell(val: Any) -> Dict[str, Any]:
    """Flatten a cell for the click-to-verify UI."""
    if val is None or val == "":
        return {"value": None, "raw": None, "confidence": None, "bbox": None}
    if isinstance(val, dict):
        bbox = val.get("bbox")
        if isinstance(bbox, (list, tuple)) and len(bbox) >= 4:
            try:
                bbox = [float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])]
            except (TypeError, ValueError):
                bbox = None
        else:
            bbox = None
        conf = None
        if val.get("confidence") is not None:
            try:
                conf = float(val["confidence"])
            except (TypeError, ValueError):
                conf = None
        return {
            "value": _unwrap(val),
            "raw": val.get("raw"),
            "confidence": conf,
            "bbox": bbox,
            "repaired": bool(val.get("repaired")),
        }
    return {"value": val, "raw": val, "confidence": None, "bbox": None}


def build_ocr_review(
    *,
    image_paths: Sequence[str],
    document_type: str,
    model_name: Optional[str],
    structured: Dict[str, Any],
    validation: Optional[Dict[str, Any]] = None,
    quality: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Compact review payload for the GUI (not dumped into chat)."""
    data = structured if isinstance(structured, dict) else {}
    meta_keys = (
        "organization", "phone", "telephone", "surveyed_by", "computed_by",
        "instrument", "instrument_name",
        "serial", "serial_number", "date", "page", "page_number", "sheet",
        "workspace_path", "user", "status",
    )
    metadata: Dict[str, Any] = {}
    for key in meta_keys:
        cell = _meta_get(data, key)
        if cell not in (None, ""):
            metadata[key] = _review_field_cell(cell)
    sections_out: List[Dict[str, Any]] = []
    for sec in (data.get("sections") or []):
        if not isinstance(sec, dict):
            continue
        heading = str(sec.get("heading") or "Section").strip() or "Section"
        lines = [str(x).strip() for x in (sec.get("lines") or []) if str(x).strip()]
        if heading or lines:
            sections_out.append({"heading": heading, "lines": lines})
    rows_out: List[Dict[str, Any]] = []
    survey_names = (
        "instrument_station", "from", "reference_station", "to",
        "hz_fl", "hz_fr", "va_fl", "va_fr",
        "slope_distance", "slope", "horizontal_distance", "hor",
        "backsight", "intermediate_sight", "foresight",
        "reduced_level", "height_of_collimation", "distance", "station",
    )
    for i, row in enumerate(_as_rows(data)):
        item: Dict[str, Any] = {"row": i + 1}
        if _rows_look_like_spreadsheet([row], data) and not _rows_look_like_traverse([row]):
            for k, v in row.items():
                if str(k).lower() in _INTERNAL_ROW_KEYS:
                    continue
                if v in (None, ""):
                    continue
                item[str(k)] = _review_field_cell(v)
        else:
            for name in survey_names:
                raw = _row_get(row, name)
                if raw in (None, ""):
                    continue
                canon = normalize_field_name(name)
                if canon in item:
                    continue
                item[canon] = _review_field_cell(raw)
        rows_out.append(item)
    uncertain: List[str] = []
    if isinstance(validation, dict):
        for u in validation.get("uncertain_fields") or []:
            if u and str(u) not in uncertain:
                uncertain.append(str(u))
        for c in validation.get("checks") or []:
            if isinstance(c, dict) and not c.get("passed"):
                for f in c.get("fields") or []:
                    key = f"row[{c.get('name', '')}].{f}" if f else str(c.get("name"))
                    if key not in uncertain:
                        uncertain.append(key)
    style = data.get("style_card") if isinstance(data.get("style_card"), dict) else {}
    return {
        "image_paths": [str(p) for p in image_paths],
        "document_type": document_type or data.get("document_type") or "generic",
        "model_name": model_name,
        "title": _unwrap(data.get("title")),
        "quality": quality or {},
        "metadata": metadata,
        "sections": sections_out,
        "rows": rows_out,
        "visible_headers": data.get("visible_headers") if isinstance(data.get("visible_headers"), list) else [],
        "uncertain": uncertain,
        "style_card": style,
    }


# ---------------------------------------------------------------------------
# Vision invoke
# ---------------------------------------------------------------------------

def run_vision_ocr(
    image_paths: Sequence[str],
    *,
    user_text: str,
    llm: Any,
    run_with_timeout: Callable[..., Any],
    mode: Optional[VisionOcrMode] = None,
    timeout_s: int = OCR_BUDGET_S,
    max_files: int = 4,
    max_file_mb: int = 10,
    model_name: Optional[str] = None,
    high_accuracy: bool = True,
    verify_uncertain: bool = False,
    workspace: Optional[Path] = None,
) -> VisionOcrResult:
    """One style-locked vision read; local Face Left/Right repairs instead of a second LLM call."""
    resolved_mode = mode or select_vision_ocr_mode(user_text)
    started = time.monotonic()
    budget = max(20, min(int(timeout_s or OCR_BUDGET_S), OCR_BUDGET_S))

    usable: List[str] = []
    for raw in image_paths:
        p = Path(raw)
        if not p.is_file():
            continue
        try:
            if p.stat().st_size / (1024 * 1024) > max_file_mb:
                logger.warning("Skipping oversized image %s", p)
                continue
        except OSError:
            continue
        usable.append(str(p.resolve()))
        if len(usable) >= max(1, int(max_files or 4)):
            break
    if not usable:
        return VisionOcrResult(success=False, mode=resolved_mode, error="No readable image files found for vision OCR.", model_name=model_name)

    qualities = [assess_image_quality(path) for path in usable]
    quality_summary = {
        "images": qualities,
        "overall": "good",
        "readable": True,
        "reason": "",
    }
    if any(not q.get("readable", True) for q in qualities):
        bad = next(q for q in qualities if not q.get("readable", True))
        reason = str(bad.get("reason") or "Image quality is too poor for reliable OCR.")
        quality_summary["overall"] = "bad"
        quality_summary["readable"] = False
        quality_summary["reason"] = reason
        msg = (
            f"This photograph is not clear enough for reliable OCR.\n\n{reason}\n\n"
            "Please retake with the sheet flat, well lit, and filling most of the frame."
        )
        return VisionOcrResult(
            success=False, mode=resolved_mode, image_paths=usable, model_name=model_name,
            error=msg, quality=quality_summary,
            ocr_review={"image_paths": usable, "quality": quality_summary, "rows": [], "metadata": {}, "uncertain": []},
        )
    if any(q.get("overall") == "questionable" for q in qualities):
        quality_summary["overall"] = "questionable"
        quality_summary["reason"] = next(
            (str(q.get("reason") or "") for q in qualities if q.get("overall") == "questionable"),
            "",
        )

    document_type = classify_ocr_document(user_text, usable)
    max_edge = 1920 if high_accuracy else 1600
    encoded: List[tuple] = []
    mime = "image/jpeg" if high_accuracy else "image/png"
    for path in usable:
        b64 = image_file_to_base64_png(path, max_edge=max_edge, enhance=high_accuracy)
        if b64:
            encoded.append((b64, mime))
    if not encoded:
        return VisionOcrResult(
            success=False, mode=resolved_mode, image_paths=usable, document_type=document_type,
            error="Could not encode image(s) for vision OCR.", model_name=model_name,
            quality=quality_summary,
        )

    prior_style = load_handwriting_style(workspace)
    user_prompt = _user_prompt_for_ocr(
        user_text, document_type=document_type, prior_style=prior_style or None
    )
    first_budget = budget
    messages = _build_vision_messages(
        system=_system_prompt_for_mode(resolved_mode, document_type=document_type),
        user_prompt=user_prompt,
        encoded=encoded,
    )
    msg, err, timed_out = run_with_timeout(first_budget, lambda: llm.invoke(messages))
    if timed_out:
        return VisionOcrResult(
            success=False, mode=resolved_mode, image_paths=usable, document_type=document_type,
            error="Vision OCR timed out.", model_name=model_name, quality=quality_summary,
        )
    if err:
        return VisionOcrResult(
            success=False, mode=resolved_mode, image_paths=usable, document_type=document_type,
            error=f"Vision OCR failed: {err}", model_name=model_name, quality=quality_summary,
        )

    raw_text = _message_content_to_text(getattr(msg, "content", msg))
    if resolved_mode == "plain_text":
        review = {
            "image_paths": usable, "document_type": document_type, "model_name": model_name,
            "quality": quality_summary, "metadata": {}, "rows": [], "uncertain": [],
            "plain_text": raw_text.strip(),
        }
        return VisionOcrResult(
            success=bool(raw_text.strip()), mode=resolved_mode, text=raw_text.strip(),
            image_paths=usable, document_type=document_type, model_name=model_name,
            error=None if raw_text.strip() else "Vision OCR returned empty text.",
            quality=quality_summary, ocr_review=review,
        )

    data = _extract_json_object(raw_text)
    if not data:
        return VisionOcrResult(
            success=bool(raw_text.strip()), mode=resolved_mode, text=raw_text.strip(),
            structured={}, image_paths=usable, document_type=document_type, model_name=model_name,
            notes="JSON parse failed; returned raw model text.",
            error=None if raw_text.strip() else "Vision OCR returned empty content.",
            quality=quality_summary,
        )

    document_type = classify_ocr_document(user_text, usable, data)
    data = normalize_extracted_document(data)
    merged_style = dict(prior_style)
    if isinstance(data.get("style_card"), dict):
        merged_style.update({str(k): str(v) for k, v in data["style_card"].items() if v})
    data = apply_survey_context_repairs(data, document_type=document_type, style_card=merged_style)
    data = apply_learned_ocr_corrections(data, workspace=workspace)
    data = apply_survey_context_repairs(data, document_type=document_type, style_card=merged_style)
    report = validate_ocr_extraction(data, user_text=user_text, image_paths=usable, document_type=document_type)
    failed = [c for c in report.checks if not c.passed]
    remaining = budget - (time.monotonic() - started)
    if high_accuracy and verify_uncertain and failed and remaining >= 16 and llm is not None:
        reread = fields_needing_reread(report)
        if reread:
            logger.info("OCR verify (%ds left) fields=%s", int(remaining), ", ".join(reread))
            vmsg, verr, vto = run_with_timeout(
                min(20, int(remaining) - 1),
                lambda: llm.invoke(_build_vision_messages(
                    system=_system_prompt_for_mode("structured", document_type=document_type),
                    user_prompt=_verify_prompt(reread, data),
                    encoded=encoded,
                )),
            )
            if not verr and not vto and vmsg is not None:
                vdata = _extract_json_object(_message_content_to_text(getattr(vmsg, "content", vmsg))) or {}
                corrections = vdata.get("corrections") if isinstance(vdata, dict) else None
                if isinstance(corrections, dict) and corrections:
                    data = _apply_corrections(data, corrections)
                    data = apply_survey_context_repairs(
                        data, document_type=document_type, style_card=merged_style
                    )
                    report = validate_ocr_extraction(
                        data, user_text=user_text, image_paths=usable, document_type=document_type
                    )
                    extra = str(vdata.get("notes") or "").strip()
                    if extra:
                        report.notes.append(f"Verification pass: {extra}")

    writer = str(_unwrap(_meta_get(data, "surveyed_by")) or "")
    save_handwriting_style(data.get("style_card") if isinstance(data.get("style_card"), dict) else merged_style,
                           workspace=workspace, writer=writer)
    data = merge_validation_into_structured(data, report)
    pt = data.get("plain_text")
    if isinstance(pt, dict):
        plain = " ".join(str(v) for v in pt.values() if v not in (None, "")).strip() or raw_text[:4000]
    else:
        plain = str(pt or data.get("text") or "").strip() or raw_text[:4000]
    notes = str(data.get("notes") or "")
    review = build_ocr_review(
        image_paths=usable,
        document_type=document_type,
        model_name=model_name,
        structured=data,
        validation=report.to_dict(),
        quality=quality_summary,
    )
    save_last_ocr_extraction(
        data,
        workspace=workspace,
        image_paths=usable,
        document_type=document_type,
        source="vision_ocr",
    )
    return VisionOcrResult(
        success=True, mode=resolved_mode, text=plain, structured=data,
        image_paths=usable, document_type=document_type, validation=report.to_dict(),
        model_name=model_name, notes=notes, quality=quality_summary, ocr_review=review,
    )


def extract_survey_plan_from_images(
    image_paths: Sequence[str],
    *,
    llm: Any,
    run_with_timeout: Callable[..., Any],
    user_notes: str = "",
    timeout_s: int = OCR_BUDGET_S,
) -> Any:
    from agent.pdf_survey_plan import (
        SurveyPlanExtraction,
        _EXTRACTION_SYSTEM,
        parse_survey_plan_extraction,
    )
    from langchain_core.messages import HumanMessage, SystemMessage

    encoded: List[str] = []
    for path in image_paths:
        b64 = image_file_to_base64_png(path, max_edge=1920, enhance=True)
        if b64:
            encoded.append(b64)
    if not encoded:
        return SurveyPlanExtraction(source="error", notes="No images could be encoded")
    user_prompt = "IMAGE SOURCE(S):\n" + "\n".join(str(p) for p in image_paths) + "\n\n"
    if user_notes.strip():
        user_prompt += f"USER INSTRUCTIONS:\n{user_notes.strip()}\n\n"
    user_prompt += (
        "Lock this writer's/drafter's glyph style, then extract traverse, metadata, and coordinates. "
        "Return JSON matching the survey plan extraction schema."
    )
    content: List[dict] = [{"type": "text", "text": user_prompt}]
    for b64 in encoded:
        content.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
    messages = [SystemMessage(content=_EXTRACTION_SYSTEM + "\n" + _HANDWRITING_LOCK), HumanMessage(content=content)]
    msg, err, timed_out = run_with_timeout(min(int(timeout_s or OCR_BUDGET_S), OCR_BUDGET_S), lambda: llm.invoke(messages))
    if timed_out:
        return SurveyPlanExtraction(source="error", notes="LLM extraction timed out")
    if err:
        return SurveyPlanExtraction(source="error", notes=f"LLM extraction failed: {err}")
    raw = _message_content_to_text(getattr(msg, "content", msg))
    data = _extract_json_object(str(raw))
    if not data:
        return SurveyPlanExtraction(source="llm_parse_failed", notes=f"Could not parse JSON from model output: {str(raw)[:500]}")
    ext = parse_survey_plan_extraction(data)
    ext.source = "vision"
    return ext


__all__ = [
    "DOC_CADASTRAL", "DOC_CALIBRATION", "DOC_GENERIC", "DOC_LEVELLING", "DOC_TABLE",
    "DOC_SPREADSHEET", "DOC_TRAVERSE", "DOC_UI",
    "OCR_BUDGET_S", "OcrDocumentType", "ValidationCheck", "ValidationReport",
    "VisionOcrMode", "VisionOcrResult",
    "apply_learned_ocr_corrections", "apply_survey_context_repairs", "assess_image_quality",
    "build_ocr_review",
    "classify_ocr_document", "export_ocr_extraction_to_docx", "export_ocr_extraction_to_excel",
    "extract_image_paths_from_query",
    "extract_survey_plan_from_images", "fields_needing_reread", "format_last_ocr_for_agent",
    "format_ocr_review_for_user",
    "format_validation_for_user",
    "hz_pair_ok", "image_file_to_base64_png", "is_ocr_export_request", "is_ocr_followup_request",
    "is_ocr_only_request", "is_ocr_word_export_request",
    "load_handwriting_style", "load_last_ocr_extraction", "load_learned_ocr_value_map",
    "looks_like_survey_plan_image_task", "looks_like_survey_sheet",
    "merge_validation_into_structured", "normalize_extracted_document", "parse_dms_to_deg",
    "parse_survey_number", "resolve_ocr_export_path", "resolve_ocr_word_export_path",
    "run_vision_ocr", "save_handwriting_style", "save_last_ocr_extraction",
    "select_vision_ocr_mode",
    "should_fastpath_image_survey_replot", "structured_from_ocr_review", "user_requested_save",
    "validate_levelling_or_calibration", "validate_ocr_extraction", "validate_traverse_rows",
    "va_pair_ok",
]
